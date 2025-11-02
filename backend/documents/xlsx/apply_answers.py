from __future__ import annotations

"""
Utilities for writing LLM answers back into structured Excel templates.

This module is the Excel counterpart to `backend.documents.docx.apply_answers`.
It takes the sanitized payload produced by `DocumentFiller` and:
* writes answer text into the correct workbook cells
* builds a companion DOCX with comment snippets so we avoid leaking citations
  directly into the spreadsheet

Both the Streamlit UI and automation workflows delegate to `write_excel_answers`.
"""

import os
import re
from typing import Any, Callable, Dict, List, Optional, Set

from openpyxl import load_workbook

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ..docx.comments import add_comment_to_run


# Pattern for [n] style citation markers in the answer text, allowing
# comma-separated values like "[1,2]" or "[1, 2, 3]"
_CITATION_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")


def _clean_excel_text(text: str) -> str:
    """Strip citations and collapse whitespace without deleting all content."""
    if not text:
        return ""
    original = text
    no_cits = _CITATION_RE.sub("", original)
    if not no_cits.strip():
        no_cits = original
    collapsed = re.sub(r"\s{2,}", " ", no_cits).strip()
    if not collapsed and no_cits.strip():
        collapsed = no_cits.strip()
    return collapsed


def _to_text_and_citations(ans: object) -> tuple[str, Dict[str, object]]:
    """Normalize answer objects to (text, {cit_num -> data})."""
    if isinstance(ans, dict):
        text = str(ans.get("text", ""))
        raw = ans.get("citations") or {}
        cits: Dict[str, object] = {}
        for k, v in raw.items():
            key = str(k)
            if isinstance(v, dict):
                snippet = v.get("text") or v.get("snippet") or v.get("content") or ""
                cits[key] = {"text": str(snippet), "source_file": v.get("source_file")}
            else:
                cits[key] = {"text": str(v)}
        return text, cits
    return str(ans or ""), {}


def _resolve_sheet_and_cell(
    entry: Dict[str, Any], worksheets: Dict[str, Any]
) -> tuple[Optional[Any], Optional[str], Optional[str], Optional[str]]:
    """Return (cell, sheet, address, skip_reason) guarding against missing slots."""
    sheet_name = entry.get("sheet") or entry.get("sheet_name")
    if "answer_cell" in entry:
        address = entry.get("answer_cell")
    else:
        address = entry.get("question_cell") or entry.get("cell") or entry.get("address")

    if sheet_name and entry.get("answer_cell") is None:
        question = (entry.get("question_text") or "").strip()
        return None, sheet_name, address, (
            f"Skipping answer for '{question}' on sheet '{sheet_name}': no answer slot"
        )
    if not sheet_name or not address:
        return None, sheet_name, address, "Skipping entry: missing sheet or cell address."

    ws = worksheets.get(sheet_name)
    if ws is None:
        return None, sheet_name, address, f"Skipping entry: sheet '{sheet_name}' not found."

    try:
        cell = ws[address]
    except Exception:
        return None, sheet_name, address, (
            f"Skipping entry: cell '{address}' not found on sheet '{sheet_name}'."
        )
    return cell, sheet_name, address, None


def _resolve_answer_payload(
    entry: Dict[str, Any],
    answer: Optional[object],
    generator: Optional[Callable[..., object]],
) -> tuple[str, str, Dict[str, object]]:
    """Normalise the answer input and optionally invoke the generator fallback."""
    question = (entry.get("question_text") or "").strip()
    current = answer
    attempt = 0

    while True:
        if current is None and generator is not None:
            print(f"DEBUG: Generating answer for question '{question}'")
            current = generator(question)
        text, citations = _to_text_and_citations(current)
        excel_text = _clean_excel_text(text)
        if excel_text or generator is None or attempt >= 1:
            if not excel_text and generator is None:
                print(f"DEBUG: No text provided for question '{question}'")
            return text, excel_text, citations
        print(
            f"DEBUG: Regenerating answer for question '{question}' due to empty text"
        )
        current = generator(question) if generator is not None else current
        attempt += 1


def _apply_text_to_cell(
    cell: Any,
    sheet_name: str,
    address: str,
    excel_text: str,
    *,
    mode: str,
) -> None:
    """Write the formatted answer into the target cell respecting the fill mode."""
    if mode == "replace":
        print(
            f"DEBUG: Replacing cell {sheet_name}!{address} contents with '{excel_text}'"
        )
        cell.value = excel_text
    elif mode == "append":
        prior = cell.value or ""
        print(
            f"DEBUG: Appending to cell {sheet_name}!{address}: prior='{prior}', new='{excel_text}'"
        )
        separator = "\n" if prior and excel_text else ""
        cell.value = f"{prior}{separator}{excel_text}"
    else:
        existing = cell.value
        if existing is None or str(existing).strip() == "":
            print(
                f"DEBUG: Filling empty cell {sheet_name}!{address} with '{excel_text}'"
            )
            cell.value = excel_text
        else:
            print(
                f"DEBUG: Cell {sheet_name}!{address} already has data; appending '{excel_text}'"
            )
            cell.value = f"{existing}\n{excel_text}"
    try:
        cell.alignment = cell.alignment.copy(wrap_text=True)
    except Exception:
        pass


def _collect_doc_entry(
    idx: int,
    entry: Dict[str, Any],
    text: str,
    citations: Dict[str, object],
) -> Optional[Dict[str, Any]]:
    """Package a single answer for inclusion in the standalone comments DOCX."""
    if not citations:
        return None
    print(f"DEBUG: Collected {len(citations)} citations for question index {idx}")
    return {
        "question": entry.get("question_text") or "",
        "text": text,
        "citations": citations,
    }


def _prepare_comments_path(
    out_xlsx_path: str,
    comments_docx_path: Optional[str],
    include_comments: bool,
) -> Optional[str]:
    """Determine where the comments DOCX should live (if at all)."""
    if not include_comments:
        return None
    if comments_docx_path:
        return comments_docx_path
    base, _ = os.path.splitext(out_xlsx_path)
    return base + "_comments.docx"


def _create_comments_document() -> docx.document.Document:
    """Create a fresh DOCX with TOC scaffolding ready for answer comment sections."""
    document = docx.Document()
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    document.settings._element.append(update)
    document.add_paragraph("Table of Contents", style="Title")
    _add_toc_field(document)
    document.add_page_break()
    return document


def _add_toc_field(document: docx.document.Document) -> None:
    """Insert a Word field pointing at the table-of-contents placeholder."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-1" \\h \\z \\u')
    run._r.append(field)


def _append_comment_section(
    document: docx.document.Document,
    entry: Dict[str, Any],
    index: int,
    total_entries: int,
) -> None:
    """Append a heading + answer paragraph (with citation comments) to the DOCX."""
    question = entry.get("question") or ""
    text = entry.get("text") or ""
    citations = entry.get("citations") or {}

    if question:
        heading = document.add_paragraph(style="Heading 1")
        lead = heading.add_run(f"Question {index}: ")
        lead.bold = True
        heading.add_run(question)

    paragraph = document.add_paragraph()
    answer_label = paragraph.add_run("Answer: ")
    answer_label.bold = True
    _add_answer_runs(document, paragraph, text, citations)

    if index < total_entries:
        document.add_page_break()


def _add_answer_runs(
    document: docx.document.Document,
    paragraph: Any,
    text: str,
    citations: Dict[str, object],
) -> None:
    """Render the answer text and attach inline comments for each citation marker."""
    position = 0
    for match in _CITATION_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        numbers = [n.strip() for n in match.group(1).split(",")]
        for offset, number in enumerate(numbers):
            run = paragraph.add_run(f"[{number}]")
            data = (
                citations.get(number)
                or citations.get(int(number))
                or citations.get(str(number))
            )
            snippet = None
            source_file = None
            if isinstance(data, dict):
                snippet = (
                    data.get("text") or data.get("snippet") or data.get("content")
                )
                source_file = data.get("source_file")
            elif data is not None:
                snippet = str(data)
            if snippet:
                add_comment_to_run(
                    document,
                    run,
                    str(snippet),
                    bold_prefix="Source Text: ",
                    source_file=source_file,
                )
            if offset < len(numbers) - 1:
                paragraph.add_run(" ")
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _write_comments_doc(entries: List[Dict[str, Any]], path: Optional[str]) -> None:
    """Persist the standalone comments DOCX, swallowing failures for older clients."""
    if not entries or not path:
        return
    try:
        document = _create_comments_document()
        total = len(entries)
        for idx, entry in enumerate(entries, start=1):
            _append_comment_section(document, entry, idx, total)
        document.save(path)
    except Exception:
        pass


def write_excel_answers(
    schema: List[Dict[str, Any]],
    answers: List[object],
    source_xlsx_path: str,
    out_xlsx_path: str,
    *,
    mode: str = "fill",  # "fill" | "replace" | "append"
    generator: Optional[Callable[..., object]] = None,
    include_comments: Optional[bool] = None,  # defaults to env RFP_INCLUDE_COMMENTS
    comments_docx_path: Optional[str] = None,
) -> Dict[str, int]:
    """
    Drop‑in replacement for the old CLI helper. Applies answers into the Excel file.

    schema[i] is expected to include:
      - 'sheet' (or 'sheet_name')
      - 'answer_cell' (if omitted we fall back to 'question_cell'; if set to
        ``None`` the answer will be skipped)
      - 'question_text' (used only if we need to generate a missing answer)

    answers[i] can be a string or a dict like:
      {"text": "...", "citations": {1: "snippet", 2: "snippet", ...}}
    Citation snippets are written to a separate Word document instead of Excel
    cell comments. If ``comments_docx_path`` is ``None``, the DOCX will be
    created next to ``out_xlsx_path`` using the same base name with a
    ``"_comments.docx"`` suffix.
    """
    inc_comments = (
        (os.getenv("RFP_INCLUDE_COMMENTS", "1") == "1")
        if include_comments is None
        else bool(include_comments)
    )

    wb = load_workbook(source_xlsx_path)
    ws_by_name = {ws.title: ws for ws in wb.worksheets}

    print(
        "DEBUG: Starting write_excel_answers with",
        len(schema),
        "schema entries and",
        len(answers),
        "answers",
    )
    print(
        f"DEBUG: Loaded workbook '{source_xlsx_path}' with {len(wb.worksheets)} worksheets"
    )

    applied = 0
    skipped = 0

    doc_entries: List[Dict[str, Any]] = []
    comments_docx_path = _prepare_comments_path(
        out_xlsx_path, comments_docx_path, inc_comments
    )

    if len(answers) < len(schema):
        answers = answers + [None] * (len(schema) - len(answers))

    for idx, entry in enumerate(schema):
        preview_sheet = entry.get("sheet") or entry.get("sheet_name")
        if "answer_cell" in entry:
            preview_addr = entry.get("answer_cell")
        else:
            preview_addr = entry.get("question_cell") or entry.get("cell") or entry.get("address")
        print(
            f"DEBUG: Processing entry {idx}: sheet='{preview_sheet}', address='{preview_addr}'"
        )

        cell, sheet_name, address, skip_reason = _resolve_sheet_and_cell(entry, ws_by_name)
        if skip_reason:
            print(skip_reason)
            skipped += 1
            continue

        text, excel_text, citations = _resolve_answer_payload(
            entry, answers[idx], generator
        )
        _apply_text_to_cell(cell, sheet_name or "", address or "", excel_text, mode=mode)

        if inc_comments:
            doc_entry = _collect_doc_entry(idx, entry, text, citations)
            if doc_entry:
                doc_entries.append(doc_entry)

        applied += 1

    wb.save(out_xlsx_path)
    wb.close()
    print(f"DEBUG: Workbook saved to {out_xlsx_path}")

    _write_comments_doc(doc_entries, comments_docx_path)

    result = {"applied": applied, "skipped": skipped, "total": len(schema)}
    print(
        "DEBUG: Finished write_excel_answers with",
        result,
    )
    return result


__all__ = ["write_excel_answers"]

# For a quick smoke test, point to a template and uncomment to run:
#     if __name__ == "__main__":
#         from pathlib import Path
#         sample_schema = [
#             {"sheet": "Responses", "answer_cell": "B2", "question_text": "Example?"}
#         ]
#         sample_answers = [{"text": "Sample answer", "citations": {1: {"text": "Snippet"}}}]
#         write_excel_answers(
#             sample_schema,
#             sample_answers,
#             source_xlsx_path="samples/template.xlsx",
#             out_xlsx_path="samples/template_filled.xlsx",
#             include_comments=True,
#         )
