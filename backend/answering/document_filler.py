from __future__ import annotations

"""
Utilities for turning QA responses into downloadable Office artifacts.

`DocumentFiller` sits between the answering pipeline and the Streamlit UI,
bridging internal answer formats with the Word/Excel writers in
`backend.documents`. The helpers here consolidate temp-file handling and the
payload shape expected by the front end.
"""

import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

# Directly use python-docx for the summary builder path.
from docx import Document

from backend.documents.docx.apply_answers import apply_answers_to_docx
from backend.documents.xlsx.apply_answers import write_excel_answers


class DocumentFiller:
    """Convert QA batches into the download payloads consumed by the Streamlit frontend."""

    def __init__(self) -> None:
        self._last_details: Dict[str, Any] = {}

    @property
    def last_details(self) -> Dict[str, Any]:
        """Expose the most recent bundle metadata for debugging/telemetry."""
        return self._last_details

    def build_excel_bundle(
        self,
        *,
        source_path: str,
        schema: List[Dict[str, Any]],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
        mode: str = "fill",
    ) -> Dict[str, Any]:
        """Return download metadata for writing an answered Excel workbook plus optional comments doc."""
        # Sanitise the QA payload so downstream writers receive only stable keys.
        answers_payload = [self._prepare_answer_for_storage(entry) for entry in qa_results]

        # Write to a temp path so Streamlit can offer a real file download.
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        tmp.close()
        write_excel_answers(
            schema,
            answers_payload,
            source_path,
            tmp.name,
            mode=mode,
            include_comments=include_citations,
        )

        downloads = [
            self._build_download(
                key="excel_answer",
                label="Download answered workbook",
                path=tmp.name,
                file_name=f"{Path(source_path).stem}_answered.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                order=10,
            )
        ]

        base, _ = os.path.splitext(tmp.name)
        comments_path = base + "_comments.docx"
        if include_citations and os.path.exists(comments_path):
            downloads.append(
                self._build_download(
                    key="excel_comments",
                    label="Download comments DOCX",
                    path=comments_path,
                    file_name=f"{Path(source_path).stem}_comments.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    order=20,
                )
            )

        qa_pairs = []
        for entry, qa in zip(schema, qa_results):
            question_text = (entry.get("question_text") or "").strip()
            qa_pairs.append(
                {
                    "question": question_text,
                    "answer": self._prepare_answer_for_display(qa, include_citations),
                    "comments": qa.get("raw_comments") or [],
                }
            )

        # Cache details so the frontend can surface the last generated bundle.
        self._last_details = {
            "mode": "excel",
            "qa_pairs": qa_pairs,
            "schema_length": len(schema),
            "downloads": downloads,
        }
        return {"downloads": downloads, "qa_pairs": qa_pairs}

    def build_docx_slot_bundle(
        self,
        *,
        source_path: str,
        slots_payload: Dict[str, Any],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
        write_mode: str = "fill",
    ) -> Dict[str, Any]:
        """Produce a DOCX file based on slot metadata plus answer payloads returned by the responder."""
        # slot_map feeds the on-disk structure consumed by the docx writer.
        slot_map: Dict[str, Dict[str, Any]] = {}
        slot_answers: Dict[str, Dict[str, Any]] = {}
        for qa in qa_results:
            slot_id = qa.get("slot_id")
            if slot_id:
                # Persist the sanitized answer so we can merge with the original
                # slot payload when writing the final document.
                storage = self._prepare_answer_for_storage(qa)
                slot_map[str(slot_id)] = storage
                slot_answers[str(slot_id)] = qa

        raw_slots = slots_payload.get("slots", []) or []
        slots_with_answers = [dict(slot) for slot in raw_slots]
        for slot in slots_with_answers:
            sid = str(slot.get("id", "") or "")
            stored_answer = slot_map.get(sid)
            if stored_answer is not None:
                slot["answer"] = stored_answer

        out_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        out_tmp.close()

        apply_answers_to_docx(
            docx_source=Path(source_path).read_bytes(),
            slots=slots_with_answers,
            mode=write_mode,
            output_path=out_tmp.name,
        )

        downloads = [
            self._build_download(
                key="docx_answer",
                label="Download answered DOCX",
                path=out_tmp.name,
                file_name=f"{Path(source_path).stem}_answered.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                order=10,
            )
        ]

        qa_pairs = []
        for slot in slots_payload.get("slots", []):
            slot_id = str(slot.get("id"))
            question_text = (slot.get("question_text") or "").strip()
            stored_answer = slot_map.get(slot_id)
            qa_entry = slot_answers.get(slot_id, {})
            qa_pairs.append(
                {
                    "question": question_text,
                    "answer": self._prepare_answer_for_display(qa_entry, include_citations),
                    "comments": qa_entry.get("raw_comments") or [],
                }
            )

        # Mirror the Excel bundle structure so downstream consumers can reuse UI components.
        self._last_details = {
            "mode": "docx_slots",
            "qa_pairs": qa_pairs,
            "downloads": downloads,
            "skipped_slots": slots_payload.get("skipped_slots") or [],
            "heuristic_skips": slots_payload.get("heuristic_skips") or [],
        }
        return {
            "downloads": downloads,
            "qa_pairs": qa_pairs,
            "skipped_slots": slots_payload.get("skipped_slots") or [],
            "heuristic_skips": slots_payload.get("heuristic_skips") or [],
        }

    def build_summary_bundle(
        self,
        *,
        questions: Sequence[str],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
    ) -> Dict[str, Any]:
        """Build a lightweight Word report summarizing free-form Q/A responses."""
        answers_text = [qa.get("answer", "") for qa in qa_results]
        comments = [qa.get("raw_comments") or [] for qa in qa_results]
        # Summaries mimic the docx layout used in slot-filling so the UX stays
        # consistent across document types.
        doc_bytes = self._generate_summary_doc(
            questions=list(questions),
            answers=answers_text,
            comments=comments,
            include_citations=include_citations,
        )

        downloads = [
            {
                "key": "document_summary",
                "label": "Download Q/A report",
                "data": doc_bytes,
                "file_name": "rfp_summary.docx",
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "order": 10,
            }
        ]

        qa_pairs = []
        for q, qa in zip(questions, qa_results):
            qa_pairs.append(
                {
                    "question": q,
                    "answer": self._prepare_answer_for_display(qa, include_citations),
                    "comments": qa.get("raw_comments") or [],
                }
            )

        # Reuse the bundled payload contract so Streamlit download widgets stay uniform.
        self._last_details = {
            "mode": "summary",
            "qa_pairs": qa_pairs,
            "downloads": downloads,
        }
        return {"downloads": downloads, "qa_pairs": qa_pairs}

    def _prepare_answer_for_storage(self, entry: Dict[str, Any]) -> Dict[str, Any]:
        """Trim the full QA payload to the minimal fields persisted to disk."""
        # Only persist the bare minimum for downstream writers; the UI keeps the
        # richer payload in-memory.
        return {
            "text": entry.get("answer", ""),
            "citations": entry.get("citations") or {},
        }

    def _prepare_answer_for_display(self, entry: Dict[str, Any], include_citations: bool) -> Any:
        """Return either raw answer text or the richer structure expected by export views."""
        text = entry.get("answer", "")
        citations = entry.get("citations") or {}
        if not include_citations:
            # Callers may strip citations for exports intended for redacted sharing.
            return text
        return {"text": text, "citations": citations}

    def _generate_summary_doc(
        self,
        *,
        questions: List[str],
        answers: List[str],
        comments: List[List],
        include_citations: bool,
    ) -> bytes:
        """Render an in-memory DOCX report that mirrors the slot-filling layout."""
        doc = Document()
        doc.add_heading("Q/A Summary", level=1)

        for idx, (question, answer, comment_list) in enumerate(zip(questions, answers, comments), start=1):
            doc.add_heading(f"Q{idx}: {question}", level=2)
            doc.add_paragraph(answer or "_No answer provided._")

            if include_citations and comment_list:
                doc.add_heading("Citations", level=3)
                for entry in comment_list:
                    try:
                        label, source, snippet, score, date = entry
                    except ValueError:
                        label, source, snippet, score, date = ("", "", str(entry), "", "")
                    bullet = []
                    if label:
                        bullet.append(f"[{label}]")
                    if source:
                        bullet.append(source)
                    if date:
                        bullet.append(str(date))
                    header = " ".join(bullet).strip() or "Source"
                    para = doc.add_paragraph(style="List Bullet")
                    para.add_run(f"{header}: ").bold = True
                    para.add_run(snippet or "No snippet provided.")

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        doc.save(tmp.name)
        data = Path(tmp.name).read_bytes()
        try:
            os.unlink(tmp.name)
        except Exception:
            pass
        return data

    def _build_download(
        self,
        *,
        key: str,
        label: str,
        path: str,
        file_name: str,
        mime: Optional[str],
        order: int,
    ) -> Dict[str, Any]:
        """Read the generated file into memory and return the format used by Streamlit downloads."""
        # Read the file into memory so we can hand Streamlit a bytes payload.
        data = Path(path).read_bytes()
        try:
            # Remove the temp artifact now that the bytes are buffered.
            os.unlink(path)
        except Exception:
            pass
        return {
            "key": key,
            "label": label,
            "data": data,
            "file_name": file_name,
            "mime": mime,
            "order": order,
        }


# To smoke-test this module without the Streamlit app, uncomment the block below.
# It wires up synthetic QA results and writes outputs to /tmp just like the UI would.
#
# if __name__ == "__main__":
#     import json
#
#     filler = DocumentFiller()
#     fake_schema = [
#         {"question_text": "Describe the fund objective.", "slot_id": 1},
#         {"question_text": "List key risks.", "slot_id": 2},
#     ]
#     fake_answers = [
#         {"answer": "The fund targets long-term capital appreciation.", "citations": {"1": {"text": "Prospectus", "source_file": "fund.pdf"}}},
#         {"answer": "Primary risks include market volatility and liquidity constraints.", "citations": {}},
#     ]
#     excel_bundle = filler.build_excel_bundle(
#         # Point to an existing workbook template in your environment.
#         source_path="samples/empty_template.xlsx",
#         schema=fake_schema,
#         qa_results=fake_answers,
#         include_citations=True,
#     )
#     print("Excel downloads:", [d["file_name"] for d in excel_bundle["downloads"]])
#     summary_bundle = filler.build_summary_bundle(
#         questions=[row["question_text"] for row in fake_schema],
#         qa_results=fake_answers,
#         include_citations=True,
#     )
#     print("Summary bytes:", len(summary_bundle["downloads"][0]["data"]))
