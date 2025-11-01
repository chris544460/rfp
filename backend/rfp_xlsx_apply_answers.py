from __future__ import annotations

import os
import re
from typing import Any, Callable, Dict, List, Optional

from openpyxl import load_workbook

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from word_comments import add_comment_to_run


# Pattern for [n] style citation markers in the answer text, allowing
# comma-separated values like "[1,2]" or "[1, 2, 3]"
_CITATION_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")


def _clean_excel_text(text: str) -> str:
    """Strip citations and collapse whitespace without deleting all content."""
    if not text:
        return ""
    original = text
    no_cits = _CITATION_RE.sub("", original)
    if not no_cits.strip():
        no_cits = original
    collapsed = re.sub(r"\s{2,}", " ", no_cits).strip()
    if not collapsed and no_cits.strip():
        collapsed = no_cits.strip()
    return collapsed


def _to_text_and_citations(ans: object) -> tuple[str, Dict[str, object]]:
    """Normalize answer objects to (text, {cit_num -> data})."""
    if isinstance(ans, dict):
        text = str(ans.get("text", ""))
        raw = ans.get("citations") or {}
        cits: Dict[str, object] = {}
        for k, v in raw.items():
            key = str(k)
            if isinstance(v, dict):
                snippet = v.get("text") or v.get("snippet") or v.get("content") or ""
                cits[key] = {"text": str(snippet), "source_file": v.get("source_file")}
            else:
                cits[key] = {"text": str(v)}
        return text, cits
    return str(ans or ""), {}


def write_excel_answers(
    schema: List[Dict[str, Any]],
    answers: List[object],
    source_xlsx_path: str,
    out_xlsx_path: str,
    *,
    mode: str = "fill",  # "fill" | "replace" | "append"
    generator: Optional[Callable[..., object]] = None,
    include_comments: Optional[bool] = None,  # defaults to env RFP_INCLUDE_COMMENTS
    comments_docx_path: Optional[str] = None,
) -> Dict[str, int]:
    """
    Drop‑in replacement for the old CLI helper. Applies answers into the Excel file.

    schema[i] is expected to include:
      - 'sheet' (or 'sheet_name')
      - 'answer_cell' (if omitted we fall back to 'question_cell'; if set to
        ``None`` the answer will be skipped)
      - 'question_text' (used only if we need to generate a missing answer)

    answers[i] can be a string or a dict like:
      {"text": "...", "citations": {1: "snippet", 2: "snippet", ...}}
    Citation snippets are written to a separate Word document instead of Excel
    cell comments. If ``comments_docx_path`` is ``None``, the DOCX will be
    created next to ``out_xlsx_path`` using the same base name with a
    ``"_comments.docx"`` suffix.
    """
    inc_comments = (
        (os.getenv("RFP_INCLUDE_COMMENTS", "1") == "1")
        if include_comments is None
        else bool(include_comments)
    )

    wb = load_workbook(source_xlsx_path)
    ws_by_name = {ws.title: ws for ws in wb.worksheets}

    print(
        "DEBUG: Starting write_excel_answers with",
        len(schema),
        "schema entries and",
        len(answers),
        "answers",
    )
    print(
        f"DEBUG: Loaded workbook '{source_xlsx_path}' with {len(wb.worksheets)} worksheets"
    )

    applied = 0
    skipped = 0

    doc_entries: List[Dict[str, Any]] = []

    if inc_comments and not comments_docx_path:
        base, _ = os.path.splitext(out_xlsx_path)
        comments_docx_path = base + "_comments.docx"

    # Ensure answers aligns with schema (allow None entries and generate if a generator is provided)
    if len(answers) < len(schema):
        answers = answers + [None] * (len(schema) - len(answers))

    for idx, ent in enumerate(schema):
        sheet_name = ent.get("sheet") or ent.get("sheet_name")
        if "answer_cell" in ent:
            addr = ent.get("answer_cell")
        else:
            addr = ent.get("question_cell") or ent.get("cell") or ent.get("address")

        print(
            f"DEBUG: Processing entry {idx}: sheet='{sheet_name}', address='{addr}'"
        )

        if not sheet_name or not addr:
            if sheet_name and ent.get("answer_cell") is None:
                qtxt = (ent.get("question_text") or "").strip()
                print(
                    f"Skipping answer for '{qtxt}' on sheet '{sheet_name}': no answer slot"
                )
            skipped += 1
            continue

        ws = ws_by_name.get(sheet_name) or wb.active
        cell = ws[addr]

        ans = answers[idx]
        if ans is None and generator:
            q = (ent.get("question_text") or "").strip()
            print(f"DEBUG: Generating answer for question '{q}'")
            ans = generator(q)

        text, citations = _to_text_and_citations(ans)
        excel_text = _clean_excel_text(text)

        if not excel_text and generator:
            q = (ent.get("question_text") or "").strip()
            print(
                f"DEBUG: Regenerating answer for question '{q}' due to empty text"
            )
            ans = generator(q)
            text, citations = _to_text_and_citations(ans)
            excel_text = _clean_excel_text(text)

        if mode == "replace":
            print(
                f"DEBUG: Replacing cell {sheet_name}!{addr} contents with '{excel_text}'"
            )
            cell.value = excel_text
        elif mode == "append":
            prior = cell.value or ""
            print(
                f"DEBUG: Appending to cell {sheet_name}!{addr}: prior='{prior}', new='{excel_text}'"
            )
            cell.value = (prior + ("\n" if prior else "") + excel_text)
        else:  # "fill" (default)
            # If cell is blank write; otherwise append on a new line to avoid clobbering
            if cell.value is None or str(cell.value).strip() == "":
                print(
                    f"DEBUG: Filling empty cell {sheet_name}!{addr} with '{excel_text}'"
                )
                cell.value = excel_text
            else:
                print(
                    f"DEBUG: Cell {sheet_name}!{addr} already has data; appending '{excel_text}'"
                )
                cell.value = f"{cell.value}\n{excel_text}"

        # Wrap long text
        try:
            cell.alignment = cell.alignment.copy(wrap_text=True)
        except Exception:
            pass

        # Collect citations for a separate Word document with real comments
        if inc_comments and citations:
            print(
                f"DEBUG: Collected {len(citations)} citations for question index {idx}"
            )
            doc_entries.append(
                {
                    "question": ent.get("question_text") or "",
                    "text": text,
                    "citations": citations,
                }
            )

        applied += 1

    wb.save(out_xlsx_path)
    wb.close()
    print(f"DEBUG: Workbook saved to {out_xlsx_path}")

    if inc_comments and comments_docx_path and doc_entries:
        try:
            doc = docx.Document()
            # Auto-populate fields (like TOC) when opened in Word
            update = OxmlElement("w:updateFields")
            update.set(qn("w:val"), "true")
            doc.settings._element.append(update)

            # First page: table of contents
            doc.add_paragraph("Table of Contents", style="Title")
            p_toc = doc.add_paragraph()
            run = p_toc.add_run()
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), 'TOC \\o "1-1" \\h \\z \\u')
            run._r.append(fld)
            doc.add_page_break()

            for idx, entry in enumerate(doc_entries, start=1):
                q = entry["question"]
                t = entry["text"]
                cits = entry["citations"]
                if q:
                    pq = doc.add_paragraph(style="Heading 1")
                    qrun = pq.add_run(f"Question {idx}: ")
                    qrun.bold = True
                    pq.add_run(q)
                pa = doc.add_paragraph()
                arun = pa.add_run("Answer: ")
                arun.bold = True
                pos = 0
                for match in _CITATION_RE.finditer(t):
                    if match.start() > pos:
                        pa.add_run(t[pos:match.start()])
                    nums = [n.strip() for n in match.group(1).split(",")]
                    for i, num in enumerate(nums):
                        run = pa.add_run(f"[{num}]")
                        data = cits.get(num) or cits.get(int(num)) or cits.get(str(num))
                        snippet = None
                        src_file = None
                        if isinstance(data, dict):
                            snippet = (
                                data.get("text")
                                or data.get("snippet")
                                or data.get("content")
                            )
                            src_file = data.get("source_file")
                        elif data is not None:
                            snippet = str(data)
                        if snippet:
                            add_comment_to_run(
                                doc, run, str(snippet), bold_prefix="Source Text: ", source_file=src_file
                            )
                        if i < len(nums) - 1:
                            pa.add_run(" ")
                    pos = match.end()
                if pos < len(t):
                    pa.add_run(t[pos:])
                if idx < len(doc_entries):
                    doc.add_page_break()
            doc.save(comments_docx_path)
        except Exception:
            pass

    result = {"applied": applied, "skipped": skipped, "total": len(schema)}
    print(
        "DEBUG: Finished write_excel_answers with",
        result,
    )
    return result


__all__ = ["write_excel_answers"]

