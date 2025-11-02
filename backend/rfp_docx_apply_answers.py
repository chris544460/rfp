#!/usr/bin/env python3
# rfp_docx_apply_answers.py
# Apply answers into a DOCX according to slots.json produced by rfp_docx_slot_finder.py

import argparse, json, os, sys, re, asyncio
import importlib
import traceback
from types import ModuleType
from typing import List, Union, Optional, Dict, Tuple, Callable

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from .rfp_docx_slot_finder import _looks_like_question

# NEW: real comment helper
from .word_comments import add_comment_to_run

# ---------------------------- Debug helpers ----------------------------
DEBUG = True
def dbg(msg: str):
    if DEBUG:
        print(f"[APPLY-DEBUG] {msg}")

# ---------------------------- DOC iteration ----------------------------
def iter_block_items(doc: docx.document.Document) -> List[Union[Paragraph, Table]]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    blocks: List[Union[Paragraph, Table]] = []
    parent = doc.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, doc))
    return blocks

def build_indexes(doc: docx.document.Document) -> Tuple[
    List[Union[Paragraph, Table]],
    List[Paragraph],
    Dict[int, int],
    Dict[int, int]
]:
    blocks = iter_block_items(doc)
    paragraphs: List[Paragraph] = []
    block_to_para: Dict[int, int] = {}
    block_to_table: Dict[int, int] = {}
    running_table_index = 0
    for bi, b in enumerate(blocks):
        if isinstance(b, Paragraph):
            block_to_para[bi] = len(paragraphs)
            paragraphs.append(b)
        elif isinstance(b, Table):
            block_to_table[bi] = running_table_index
            running_table_index += 1
    return blocks, paragraphs, block_to_para, block_to_table

# ---------------------------- Utilities ----------------------------
_BLANK_RE = re.compile(r"_+\s*$")
_CHECKBOX_CHARS = "☐☑☒□■✓✔✗✘"
# Allow comma-separated citations like "[1,2]" or "[1, 2, 3]"
_CITATION_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")

def is_blank_para(p: Paragraph) -> bool:
    t = (p.text or "").strip()
    if t == "":
        return True
    if _BLANK_RE.match(t):
        return True
    if re.fullmatch(r"\[(?:insert|enter|provide)[^\]]*\]", t.lower()):
        return True
    try:
        if any(r.text and r.underline for r in p.runs) and len(t.replace("_","").strip()) == 0:
            return True
    except Exception:
        pass
    return False

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p_elm = OxmlElement("w:p")
    paragraph._element.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)
    if text:
        new_p.add_run(text)
    return new_p

def normalize_question(q: str) -> str:
    return " ".join((q or "").strip().lower().split())

def _normalize_citations(raw: object) -> Dict[str, object]:
    if not raw:
        return {}
    result: Dict[str, object] = {}
    if isinstance(raw, dict):
        items = raw.items()
    elif isinstance(raw, list):
        items = []
        for i, item in enumerate(raw, 1):
            key = getattr(item, "get", lambda *_: None)("id") or getattr(item, "get", lambda *_: None)("num") or i
            items.append((key, item))
    else:
        return {}
    for key, val in items:
        if isinstance(val, dict):
            snippet = val.get("text") or val.get("snippet") or val.get("source_text") or val.get("content") or ""
            result[str(key)] = {"text": str(snippet), "source_file": val.get("source_file")}
        else:
            result[str(key)] = {"text": str(val)}
    return result

def _append_with_bold(paragraph: Paragraph, text: str, bold_state: bool) -> bool:
    """Append text to paragraph, interpreting **markers** as bold toggles."""
    if not text:
        return bold_state
    i = 0
    length = len(text)
    while i < length:
        if text.startswith("**", i):
            bold_state = not bold_state
            i += 2
            continue
        next_marker = text.find("**", i)
        if next_marker == -1:
            segment = text[i:]
            i = length
        else:
            segment = text[i:next_marker]
            i = next_marker
        if segment:
            run = paragraph.add_run(segment)
            if bold_state:
                run.bold = True
    return bold_state


def _add_text_with_citations(paragraph: Paragraph, text: str, citations: Dict[object, object]) -> None:
    """Write text and attach Word comments to each [n] marker using Utilities helper."""
    doc = paragraph.part.document
    parts = text.split("\n")
    bold_state = False
    for li, line in enumerate(parts):
        pos = 0
        for match in _CITATION_RE.finditer(line):
            if match.start() > pos:
                bold_state = _append_with_bold(paragraph, line[pos:match.start()], bold_state)
            nums = [n.strip() for n in match.group(1).split(",")]
            for i, num in enumerate(nums):
                run = paragraph.add_run(f"[{num}]")
                data = citations.get(num) or citations.get(int(num))
                snippet = None
                source_file = None
                if isinstance(data, dict):
                    snippet = data.get("text") or data.get("snippet") or data.get("content")
                    source_file = data.get("source_file")
                elif data is not None:
                    snippet = str(data)
                if snippet:
                    add_comment_to_run(
                        doc,
                        run,
                        str(snippet),
                        bold_prefix="Source Text: ",
                        source_file=source_file,
                    )
                if i < len(nums) - 1:
                    paragraph.add_run(" ")
            pos = match.end()
        if pos < len(line):
            bold_state = _append_with_bold(paragraph, line[pos:], bold_state)
        if li < len(parts) - 1:
            paragraph.add_run().add_break()

# ---------------------------- Answers loader ----------------------------
def load_answers(answers_path: str) -> Tuple[Dict[str, object], Dict[str, object]]:
    with open(answers_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    by_id: Dict[str, object] = {}
    by_q: Dict[str, object] = {}
    if isinstance(data, dict):
        if "by_id" in data or "by_question" in data:
            for k, v in (data.get("by_id") or {}).items():
                by_id[str(k)] = v
            for k, v in (data.get("by_question") or {}).items():
                by_q[normalize_question(k)] = v
        else:
            for k, v in data.items():
                kstr = str(k)
                if kstr.startswith("slot_"):
                    by_id[kstr] = v
                else:
                    by_q[normalize_question(kstr)] = v
    elif isinstance(data, list):
        for item in data:
            if not isinstance(item, dict):
                continue
            if "slot_id" in item:
                by_id[str(item["slot_id"])] = item.get("answer", "")
            elif "question_text" in item:
                by_q[normalize_question(str(item["question_text"]))] = item.get("answer", "")
    else:
        raise ValueError("Unsupported answers JSON structure.")
    return by_id, by_q

# ---------------------------- Locator resolution ----------------------------
def resolve_anchor_paragraph(
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    paragraphs: List[Paragraph],
    block_to_para: Dict[int, int],
    locator: Dict[str, object],
    meta: Optional[Dict[str, object]]
) -> Optional[Paragraph]:
    ltype = str(locator.get("type", ""))
    p_idx = locator.get("paragraph_index")
    p_idx = int(p_idx) if p_idx is not None else None
    if ltype == "paragraph":
        if p_idx is None:
            return None
        if 0 <= p_idx < len(paragraphs):
            return paragraphs[p_idx]
        return None
    if ltype == "paragraph_after":
        qb = None
        if isinstance(meta, dict) and "q_block" in meta:
            try:
                qb = int(meta["q_block"])
            except Exception:
                qb = None
        if qb is not None and 0 <= qb < len(blocks) and isinstance(blocks[qb], Paragraph):
            return blocks[qb]  # type: ignore
        if p_idx is not None and 0 <= p_idx < len(blocks) and isinstance(blocks[p_idx], Paragraph):
            return blocks[p_idx]  # type: ignore
        if p_idx is not None and 0 <= p_idx < len(paragraphs):
            return paragraphs[p_idx]
    return None

def get_target_paragraph_after_anchor(
    blocks: List[Union[Paragraph, Table]],
    block_to_para: Dict[int, int],
    paragraphs: List[Paragraph],
    anchor_para: Paragraph,
    offset: int
) -> Paragraph:
    anchor_para_index = None
    anchor_block_index = None
    for bi, b in enumerate(blocks):
        if isinstance(b, Paragraph) and b is anchor_para:
            anchor_block_index = bi
            anchor_para_index = block_to_para[bi]
            break
    if anchor_block_index is None or anchor_para_index is None:
        anchor_para_index = paragraphs.index(anchor_para)
    subsequent_paras: List[Paragraph] = []
    if anchor_block_index is not None:
        for b in blocks[anchor_block_index + 1:]:
            if isinstance(b, Paragraph):
                txt = (b.text or "").strip()
                if _looks_like_question(txt):
                    break
                subsequent_paras.append(b)
    else:
        for p in paragraphs[anchor_para_index + 1:]:
            txt = (p.text or "").strip()
            if _looks_like_question(txt):
                break
            subsequent_paras.append(p)
    if len(subsequent_paras) >= offset:
        return subsequent_paras[offset - 1]
    needed = offset - len(subsequent_paras)
    dbg(f"Not enough following paragraphs: need to insert {needed} after anchor")
    last = anchor_para if not subsequent_paras else subsequent_paras[-1]
    created: Paragraph = last
    for _ in range(needed):
        created = insert_paragraph_after(created, "")
        subsequent_paras.append(created)
    return subsequent_paras[offset - 1]

# ---------------------------- Apply operations ----------------------------
def apply_to_paragraph(target: Paragraph, answer: object, mode: str = "fill") -> None:
    if isinstance(answer, dict):
        answer_text = str(answer.get("text", ""))
        citations = _normalize_citations(answer.get("citations"))
    else:
        answer_text = str(answer)
        citations = {}
    existing = target.text or ""
    if mode == "replace":
        target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        return
    if mode == "append":
        if existing:
            target.add_run("\n")
        else:
            target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        return
    if is_blank_para(target) or not existing.strip():
        target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        return

    answer_norm = answer_text.strip()
    if existing.strip() == answer_norm:
        target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        dbg("Replaced existing matching answer in target paragraph.")
        return

    next_p_elm = target._p.getnext()
    next_para = None
    if next_p_elm is not None and next_p_elm.tag.endswith("p"):
        next_para = Paragraph(next_p_elm, target._parent)
    if next_para and (next_para.text or "").strip() == answer_norm:
        next_para.text = ""
        _add_text_with_citations(next_para, answer_text, citations)
        dbg("Replaced matching answer in subsequent paragraph.")
        return

    new_p = insert_paragraph_after(target, "")
    _add_text_with_citations(new_p, answer_text, citations)
    dbg("Target paragraph not blank; appended answer in a new paragraph below.")

def apply_to_table_cell(tbl: Table, row: int, col: int, answer: object, mode: str = "fill") -> None:
    try:
        cell = tbl.cell(row, col)
    except Exception:
        return
    if isinstance(answer, dict):
        answer_text = str(answer.get("text", ""))
        citations = _normalize_citations(answer.get("citations"))
    else:
        answer_text = str(answer)
        citations = {}
    current = cell.text or ""
    if mode == "replace":
        cell.text = ""
        p = cell.paragraphs[0]
        _add_text_with_citations(p, answer_text, citations)
        return
    if current.strip():
        p = cell.paragraphs[-1]
        p.add_run().add_break()
        _add_text_with_citations(p, answer_text, citations)
    else:
        cell.text = ""
        p = cell.paragraphs[0]
        _add_text_with_citations(p, answer_text, citations)

def mark_multiple_choice(
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    choices_meta: List[Dict[str, object]],
    index: int,
    style: Optional[str] = None,
    comment_text: Optional[str] = None,
) -> None:
    if not isinstance(choices_meta, list):
        return
    if index is None or not (0 <= index < len(choices_meta)):
        return
    meta = choices_meta[index]
    b_idx = int(meta.get("block_index", -1))
    if not (0 <= b_idx < len(blocks)):
        return
    para = blocks[b_idx]
    if not isinstance(para, Paragraph):
        return
    prefix = str(meta.get("prefix", ""))
    text = para.text or ""
    if style in (None, "", "auto"):
        if any(ch in prefix for ch in _CHECKBOX_CHARS):
            style = "checkbox"
        elif prefix.strip() in ("()", "[]"):
            style = "fill"
        else:
            style = "highlight"
    if style == "checkbox" and any(ch in prefix for ch in _CHECKBOX_CHARS):
        para.text = re.sub(rf"[{_CHECKBOX_CHARS}]", "☑", text, count=1)
    elif style == "fill" and prefix.strip() in ("()", "[]"):
        mark = prefix[0] + "X" + prefix[1]
        para.text = text.replace(prefix, mark, 1)
    elif style == "highlight":
        for run in para.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        para.text = "X " + text
    if comment_text:
        run = para.runs[0] if para.runs else para.add_run()
        try:
            add_comment_to_run(
                doc, run, comment_text, bold_prefix="Source Text: ", source_file=None
            )
        except Exception as e:
            dbg(f"  -> error adding comment: {e}")


def _build_answers_map(
    slots: List[Dict[str, object]],
    *,
    by_id: Dict[str, object],
    by_q: Dict[str, object],
    generator: Optional[Callable[..., object]],
    gen_name: str,
) -> Tuple[Dict[str, Optional[object]], int]:
    answers: Dict[str, Optional[object]] = {}
    to_generate: List[Tuple[str, str, dict]] = []
    for slot in slots:
        sid = slot.get("id", "")
        question_text = (slot.get("question_text") or "").strip()
        meta = slot.get("meta") or {}
        answer: Optional[object]
        if sid in by_id:
            answer = by_id[sid]
        else:
            key = normalize_question(question_text)
            answer = by_q.get(key)
        if answer is None and generator is not None:
            if not question_text:
                dbg(f"Skipping generation for slot {sid}: blank question text")
            else:
                kwargs: Dict[str, object] = {}
                if slot.get("answer_type") == "multiple_choice":
                    choice_meta = meta.get("choices", [])
                    kwargs["choices"] = [
                        c.get("text") if isinstance(c, dict) else str(c)
                        for c in choice_meta
                    ]
                    kwargs["choice_meta"] = choice_meta
                to_generate.append((sid, question_text, kwargs))
        answers[sid] = answer

    generated = 0
    if to_generate and generator is not None:

        async def run_all() -> List[Tuple[str, Optional[object]]]:
            async def worker(sid: str, question: str, kwargs: Dict[str, object]):
                try:
                    ans = await asyncio.to_thread(generator, question, **kwargs)
                    dbg(f"Generated answer via {gen_name} for slot {sid}: {ans}")
                    return sid, ans
                except Exception as exc:
                    dbg(f"Generator error for question '{question}': {exc}")
                    dbg(f"Generator error details: {traceback.format_exc()}")
                    return sid, None

            tasks = [asyncio.create_task(worker(*item)) for item in to_generate]
            return await asyncio.gather(*tasks)

        for sid, ans in asyncio.run(run_all()):
            if ans is not None:
                answers[sid] = ans
                generated += 1

    return answers, generated


def _format_citation_comment(raw: object) -> Optional[str]:
    if not isinstance(raw, dict):
        return None
    parts: List[str] = []
    for value in raw.values():
        if isinstance(value, dict):
            snippet = (
                value.get("text")
                or value.get("snippet")
                or value.get("content")
                or ""
            )
            source_file = value.get("source_file")
            piece = str(snippet)
            if source_file:
                piece += f"\nSource File:\n {source_file}"
            parts.append(piece)
        else:
            parts.append(str(value))
    return "\n\n".join(parts) if parts else None


def _apply_multiple_choice_slot(
    *,
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    slot: Dict[str, object],
    answer: Dict[str, object],
) -> str:
    choice_meta = (slot.get("meta") or {}).get("choices", [])
    if not choice_meta:
        dbg("  -> no choice metadata present")
        return "bad_locator"
    idx = answer.get("choice_index")
    if idx is None:
        dbg("  -> could not resolve selected choice index")
        return "bad_locator"
    style = answer.get("style")
    comment_text = _format_citation_comment(answer.get("citations"))
    try:
        mark_multiple_choice(
            doc,
            blocks,
            choice_meta,
            int(idx),
            style,
            comment_text,
        )
        dbg("  -> marked choice in-place")
        return "applied"
    except Exception as exc:
        dbg(f"  -> error marking multiple choice: {exc}")
        return "bad_locator"


def _apply_table_cell_slot(
    *,
    doc: docx.document.Document,
    locator: Dict[str, object],
    answer: object,
    mode: str,
) -> str:
    t_index = locator.get("table_index")
    if t_index is None:
        dbg("  -> bad locator: missing table_index")
        return "bad_locator"
    try:
        table_idx = int(t_index)
        row = int(locator.get("row", 0))
        col = int(locator.get("col", 0))
    except Exception:
        dbg("  -> invalid table coordinates")
        return "bad_locator"
    if not (0 <= table_idx < len(doc.tables)):
        dbg(
            f"  -> table_index {table_idx} out of bounds; have {len(doc.tables)} tables"
        )
        return "table_oob"
    tbl = doc.tables[table_idx]
    apply_to_table_cell(tbl, row, col, answer, mode=mode)
    dbg(f"  -> wrote into table[{table_idx}] cell({row},{col})")
    return "applied"


def _apply_paragraph_slot(
    *,
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    paragraphs: List[Paragraph],
    block_to_para: Dict[int, int],
    locator: Dict[str, object],
    meta: Dict[str, object],
    answer: object,
    mode: str,
    ltype: str,
) -> str:
    anchor = resolve_anchor_paragraph(doc, blocks, paragraphs, block_to_para, locator, meta)
    if anchor is None:
        dbg("  -> could not resolve anchor/target paragraph")
        return "bad_locator"

    force_after_question = bool(meta.get("force_insert_after_question"))
    if ltype == "paragraph_after":
        offset = int(locator.get("offset", 1) or 1)
        if force_after_question:
            target = insert_paragraph_after(anchor, "")
        else:
            target = get_target_paragraph_after_anchor(
                blocks,
                block_to_para,
                paragraphs,
                anchor,
                offset,
            )
    else:
        target = anchor

    apply_to_paragraph(target, answer, mode=mode)
    dbg("  -> wrote into paragraph")
    return "applied"


def _apply_slot_to_doc(
    slot: Dict[str, object],
    answer: object,
    *,
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    paragraphs: List[Paragraph],
    block_to_para: Dict[int, int],
    mode: str,
) -> str:
    answer_type = slot.get("answer_type")
    meta = slot.get("meta") or {}
    locator = slot.get("answer_locator") or {}
    ltype = str(locator.get("type", ""))

    if (
        answer_type == "multiple_choice"
        and isinstance(answer, dict)
        and "choice_index" in answer
    ):
        return _apply_multiple_choice_slot(
            doc=doc,
            blocks=blocks,
            slot=slot,
            answer=answer,
        )

    dbg(f"Applying answer for slot {slot.get('id', '')} (type={ltype})")

    if ltype == "table_cell":
        return _apply_table_cell_slot(
            doc=doc,
            locator=locator,
            answer=answer,
            mode=mode,
        )
    if ltype in ("paragraph_after", "paragraph"):
        return _apply_paragraph_slot(
            doc=doc,
            blocks=blocks,
            paragraphs=paragraphs,
            block_to_para=block_to_para,
            locator=locator,
            meta=meta,
            answer=answer,
            mode=mode,
            ltype=ltype,
        )
    dbg(f"  -> unsupported locator type: {ltype}")
    return "bad_locator"


# ---------------------------- Main application flow ----------------------------
def apply_answers_to_docx(
    docx_path: str,
    slots_json_path: str,
    answers_json_path: str,
    out_path: str,
    mode: str = "fill",
    generator: Optional[Callable[..., object]] = None,
    gen_name: str = ""
) -> Dict[str, int]:
    with open(slots_json_path, "r", encoding="utf-8") as f:
        slots_payload = json.load(f)

    by_id, by_q = ({}, {})
    if answers_json_path and answers_json_path != "-" and os.path.isfile(answers_json_path):
        by_id, by_q = load_answers(answers_json_path)
        dbg(f"Answers loaded: by_id={len(by_id)}, by_question={len(by_q)}")
    else:
        dbg("No answers file provided; relying solely on generator (if any)")

    doc = docx.Document(docx_path)
    blocks, paragraphs, block_to_para, block_to_table = build_indexes(doc)

    applied = 0
    skipped_no_answer = 0
    skipped_bad_locator = 0
    skipped_table_oob = 0
    generated = 0

    slots = (slots_payload or {}).get("slots", [])

    answers, generated = _build_answers_map(
        slots,
        by_id=by_id,
        by_q=by_q,
        generator=generator,
        gen_name=gen_name,
    )

    for s in slots:
        sid = s.get("id", "")
        question_text = (s.get("question_text") or "").strip()
        answer = answers.get(sid)
        if answer is None:
            dbg(f"NO ANSWER for slot {sid!r} / question '{question_text}' — skipping")
            skipped_no_answer += 1
            continue
        try:
            status = _apply_slot_to_doc(
                s,
                answer,
                doc=doc,
                blocks=blocks,
                paragraphs=paragraphs,
                block_to_para=block_to_para,
                mode=mode,
            )
        except Exception as exc:
            dbg(f"  -> error while applying slot {sid}: {exc}")
            status = "bad_locator"

        if status == "applied":
            applied += 1
        elif status == "table_oob":
            skipped_table_oob += 1
        else:
            skipped_bad_locator += 1

    doc.save(out_path)
    print(f"Wrote {out_path}")

    return {
        "applied": applied,
        "skipped_no_answer": skipped_no_answer,
        "skipped_bad_locator": skipped_bad_locator,
        "skipped_table_oob": skipped_table_oob,
        "total_slots": len(slots),
        "generated": generated,
    }

# ---------------------------- CLI ----------------------------
def main():
    ap = argparse.ArgumentParser(description="Apply answers into a DOCX using slots.json")
    ap.add_argument("docx_path", help="Path to the original .docx")
    ap.add_argument("slots_json", help="Path to slots.json produced by the detector")
    ap.add_argument("answers_json", nargs="?", default="", help="Path to answers.json (optional if using --generate)")
    ap.add_argument("-o", "--out", required=True, help="Path to write updated .docx")
    ap.add_argument("--mode", choices=["replace", "append", "fill"], default="fill",
                    help="Write mode for paragraphs/cells (default: fill)")
    ap.add_argument(
        "--debug",
        dest="debug",
        action="store_true",
        default=True,
        help="Verbose debug logging (default on)",
    )
    ap.add_argument(
        "--no-debug",
        dest="debug",
        action="store_false",
        help="Disable debug logging",
    )
    ap.add_argument("--generate", metavar="MODULE:FUNC", help="Dynamically generate answers by calling given function for each question (e.g. rfp_utils.my_module:gen_answer)")
    if len(sys.argv) == 1:
        ap.print_help()
        sys.exit(1)
    args = ap.parse_args()

    global DEBUG
    DEBUG = args.debug
    if DEBUG:
        print("### APPLY DEBUG MODE ON ###")
        print(f"[apply_answers] source={args.docx_path} slots={args.slots_json}")

    required_paths = [args.docx_path, args.slots_json]
    for p in required_paths:
        if not os.path.isfile(p):
            print(f"Error: '{p}' does not exist.", file=sys.stderr)
            sys.exit(1)
    if DEBUG:
        print("[apply_answers] validated input paths")
    if args.answers_json and args.answers_json != "-" and not os.path.isfile(args.answers_json):
        if not args.generate:
            print(f"Error: answers file '{args.answers_json}' does not exist and no --generate specified.", file=sys.stderr)
            sys.exit(1)

    gen_callable = None
    gen_name = ""
    if args.generate:
        if ":" not in args.generate:
            print("Error: --generate requires MODULE:FUNC", file=sys.stderr)
            sys.exit(1)
        mod_name, func_name = args.generate.split(":", 1)
        try:
            module: ModuleType = importlib.import_module(mod_name)
            gen_callable = getattr(module, func_name)
            if not callable(gen_callable):
                raise AttributeError
            gen_name = args.generate
            if DEBUG:
                print(f"Loaded generator function {gen_name}")
        except Exception as e:
            print(f"Error: failed to load generator function {args.generate}: {e}", file=sys.stderr)
            sys.exit(1)

    try:
        if DEBUG:
            print("[apply_answers] applying answers to document")
        summary = apply_answers_to_docx(
            args.docx_path,
            args.slots_json,
            args.answers_json,
            args.out,
            mode=args.mode,
            generator=gen_callable,
            gen_name=gen_name
        )
    except Exception as e:
        print(f"Error: failed to apply answers: {e}", file=sys.stderr)
        sys.exit(1)

    if DEBUG:
        print("--- APPLY SUMMARY ---")
        for k, v in summary.items():
            print(f"{k}: {v}")

if __name__ == "__main__":
    main()
