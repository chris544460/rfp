from __future__ import annotations

import json
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from docx import Document
from ..rfp_docx_apply_answers import apply_answers_to_docx
from ..rfp_xlsx_apply_answers import write_excel_answers


class DocumentFiller:
    """Helper that turns answer batches into downloadable artifacts."""

    def __init__(self) -> None:
        self._last_details: Dict[str, Any] = {}

    @property
    def last_details(self) -> Dict[str, Any]:
        return self._last_details

    def build_excel_bundle(
        self,
        *,
        source_path: str,
        schema: List[Dict[str, Any]],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
        mode: str = "fill",
    ) -> Dict[str, Any]:
        answers_payload = [self._prepare_answer_for_storage(entry) for entry in qa_results]

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        tmp.close()
        write_excel_answers(
            schema,
            answers_payload,
            source_path,
            tmp.name,
            mode=mode,
            include_comments=include_citations,
        )

        downloads = [
            self._build_download(
                key="excel_answer",
                label="Download answered workbook",
                path=tmp.name,
                file_name=f"{Path(source_path).stem}_answered.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                order=10,
            )
        ]

        base, _ = os.path.splitext(tmp.name)
        comments_path = base + "_comments.docx"
        if include_citations and os.path.exists(comments_path):
            downloads.append(
                self._build_download(
                    key="excel_comments",
                    label="Download comments DOCX",
                    path=comments_path,
                    file_name=f"{Path(source_path).stem}_comments.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    order=20,
                )
            )

        qa_pairs = []
        for entry, qa in zip(schema, qa_results):
            question_text = (entry.get("question_text") or "").strip()
            qa_pairs.append(
                {
                    "question": question_text,
                    "answer": self._prepare_answer_for_display(qa, include_citations),
                    "comments": qa.get("raw_comments") or [],
                }
            )

        self._last_details = {
            "mode": "excel",
            "qa_pairs": qa_pairs,
            "schema_length": len(schema),
            "downloads": downloads,
        }
        return {"downloads": downloads, "qa_pairs": qa_pairs}

    def build_docx_slot_bundle(
        self,
        *,
        source_path: str,
        slots_payload: Dict[str, Any],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
        write_mode: str = "fill",
    ) -> Dict[str, Any]:
        slot_map: Dict[str, Dict[str, Any]] = {}
        slot_answers: Dict[str, Dict[str, Any]] = {}
        for qa in qa_results:
            slot_id = qa.get("slot_id")
            if slot_id:
                storage = self._prepare_answer_for_storage(qa)
                slot_map[str(slot_id)] = storage
                slot_answers[str(slot_id)] = qa

        answers_tmp = tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", delete=False, suffix=".json")
        json.dump({"by_id": slot_map}, answers_tmp)
        answers_tmp.flush()
        answers_tmp.close()

        slots_tmp = tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", delete=False, suffix=".json")
        json.dump(slots_payload, slots_tmp)
        slots_tmp.flush()
        slots_tmp.close()

        out_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        out_tmp.close()

        apply_answers_to_docx(
            docx_path=source_path,
            slots_json_path=slots_tmp.name,
            answers_json_path=answers_tmp.name,
            out_path=out_tmp.name,
            mode=write_mode,
            generator=None,
            gen_name="document_filler.responder",
        )

        for temp_path in (answers_tmp.name, slots_tmp.name):
            try:
                os.unlink(temp_path)
            except Exception:
                pass

        downloads = [
            self._build_download(
                key="docx_answer",
                label="Download answered DOCX",
                path=out_tmp.name,
                file_name=f"{Path(source_path).stem}_answered.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                order=10,
            )
        ]

        qa_pairs = []
        for slot in slots_payload.get("slots", []):
            slot_id = str(slot.get("id"))
            question_text = (slot.get("question_text") or "").strip()
            stored_answer = slot_map.get(slot_id)
            qa_entry = slot_answers.get(slot_id, {})
            qa_pairs.append(
                {
                    "question": question_text,
                    "answer": self._prepare_answer_for_display(qa_entry, include_citations),
                    "comments": qa_entry.get("raw_comments") or [],
                }
            )

        self._last_details = {
            "mode": "docx_slots",
            "qa_pairs": qa_pairs,
            "downloads": downloads,
            "skipped_slots": slots_payload.get("skipped_slots") or [],
            "heuristic_skips": slots_payload.get("heuristic_skips") or [],
        }
        return {
            "downloads": downloads,
            "qa_pairs": qa_pairs,
            "skipped_slots": slots_payload.get("skipped_slots") or [],
            "heuristic_skips": slots_payload.get("heuristic_skips") or [],
        }

    def build_summary_bundle(
        self,
        *,
        questions: Sequence[str],
        qa_results: Sequence[Dict[str, Any]],
        include_citations: bool,
    ) -> Dict[str, Any]:
        answers_text = [qa.get("answer", "") for qa in qa_results]
        comments = [qa.get("raw_comments") or [] for qa in qa_results]
        doc_bytes = self._generate_summary_doc(
            questions=list(questions),
            answers=answers_text,
            comments=comments,
            include_citations=include_citations,
        )

        downloads = [
            {
                "key": "document_summary",
                "label": "Download Q/A report",
                "data": doc_bytes,
                "file_name": "rfp_summary.docx",
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "order": 10,
            }
        ]

        qa_pairs = []
        for q, qa in zip(questions, qa_results):
            qa_pairs.append(
                {
                    "question": q,
                    "answer": self._prepare_answer_for_display(qa, include_citations),
                    "comments": qa.get("raw_comments") or [],
                }
            )

        self._last_details = {
            "mode": "summary",
            "qa_pairs": qa_pairs,
            "downloads": downloads,
        }
        return {"downloads": downloads, "qa_pairs": qa_pairs}

    def _prepare_answer_for_storage(self, entry: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "text": entry.get("answer", ""),
            "citations": entry.get("citations") or {},
        }

    def _prepare_answer_for_display(self, entry: Dict[str, Any], include_citations: bool) -> Any:
        text = entry.get("answer", "")
        citations = entry.get("citations") or {}
        if not include_citations:
            return text
        return {"text": text, "citations": citations}

    def _generate_summary_doc(
        self,
        *,
        questions: List[str],
        answers: List[str],
        comments: List[List],
        include_citations: bool,
    ) -> bytes:
        doc = Document()
        doc.add_heading("Q/A Summary", level=1)

        for idx, (question, answer, comment_list) in enumerate(zip(questions, answers, comments), start=1):
            doc.add_heading(f"Q{idx}: {question}", level=2)
            doc.add_paragraph(answer or "_No answer provided._")

            if include_citations and comment_list:
                doc.add_heading("Citations", level=3)
                for entry in comment_list:
                    try:
                        label, source, snippet, score, date = entry
                    except ValueError:
                        label, source, snippet, score, date = ("", "", str(entry), "", "")
                    bullet = []
                    if label:
                        bullet.append(f"[{label}]")
                    if source:
                        bullet.append(source)
                    if date:
                        bullet.append(str(date))
                    header = " ".join(bullet).strip() or "Source"
                    para = doc.add_paragraph(style="List Bullet")
                    para.add_run(f"{header}: ").bold = True
                    para.add_run(snippet or "No snippet provided.")

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        doc.save(tmp.name)
        data = Path(tmp.name).read_bytes()
        try:
            os.unlink(tmp.name)
        except Exception:
            pass
        return data

    def _build_download(
        self,
        *,
        key: str,
        label: str,
        path: str,
        file_name: str,
        mime: Optional[str],
        order: int,
    ) -> Dict[str, Any]:
        data = Path(path).read_bytes()
        try:
            os.unlink(path)
        except Exception:
            pass
        return {
            "key": key,
            "label": label,
            "data": data,
            "file_name": file_name,
            "mime": mime,
            "order": order,
        }
