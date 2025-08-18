#!/usr/bin/env python3
from __future__ import annotations
# rfp_docx_slot_finder.py
#
# Requires environment variables for the chosen framework:
#   • Framework selection: ANSWER_FRAMEWORK=openai|aladdin
#   • OpenAI: set OPENAI_API_KEY (and optional OPENAI_MODEL)
#   • Aladdin: set aladdin_studio_api_key, defaultWebServer, aladdin_user, aladdin_passwd

DEBUG = True
SHOW_TEXT = False  # when True, print full prompt/completion payloads

# ────────────── Token/cost tracking and pricing map ──────────────
TOTAL_INPUT_TOKENS = 0
TOTAL_OUTPUT_TOKENS = 0
TOTAL_COST_USD = 0.0

# very approximate per‑1K‑token costs in USD (Aug 2025 public pricing)
# GPT-5 nano: the fastest, cheapest version of GPT-5—great for summarization and classification tasks
MODEL_PRICING = {
    # $0.05 input / $0.005 cached input / $0.40 output per million tokens
    "gpt-5-nano": {"in": 0.00005, "out": 0.0004, "cached_in": 0.000005},
    "gpt-4o": {"in": 0.005, "out": 0.015},          # $5 / $15 per million
    "gpt-4o-mini": {"in": 0.003, "out": 0.009},     # hypothetical mini tier
    "gpt-4o-max": {"in": 0.007, "out": 0.021},      # hypothetical tier
}

def _record_usage(model: str, usage: Dict[str, int]):
    """Accumulate token usage and cost into globals."""
    global TOTAL_INPUT_TOKENS, TOTAL_OUTPUT_TOKENS, TOTAL_COST_USD
    prompt_toks = usage.get("prompt_tokens", 0)
    compl_toks = usage.get("completion_tokens", 0)
    TOTAL_INPUT_TOKENS += prompt_toks
    TOTAL_OUTPUT_TOKENS += compl_toks
    price = MODEL_PRICING.get(model, MODEL_PRICING.get("gpt-5-nano"))
    cost = (prompt_toks / 1000) * price["in"] + (compl_toks / 1000) * price["out"]
    TOTAL_COST_USD += cost
    if DEBUG:
        dbg(f"Cost for call [{model}]: input {prompt_toks} tok, output {compl_toks} tok → ${cost:.6f}")

def dbg(msg: str):
    if DEBUG:
        print(f"[DEBUG] {msg}")
import os, re, json, uuid, argparse
import sys
import math
import asyncio
from dataclasses import dataclass, asdict
from typing import List, Optional, Dict, Any, Tuple, Union, Set
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv

# Load environment variables from a .env file if present
load_dotenv(override=True)

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from answer_composer import CompletionsClient, get_openai_completion
from prompts import read_prompt

# Framework and model selection
FRAMEWORK = os.getenv("ANSWER_FRAMEWORK", "openai")
MODEL = os.getenv("OPENAI_MODEL", "gpt-5-nano")


def _call_llm(prompt: str, json_output: bool = False) -> str:
    """Call the selected LLM framework and record usage."""
    if FRAMEWORK == "aladdin":
        client = CompletionsClient(model=MODEL)
        content, usage = client.get_completion(prompt, json_output=json_output)
    else:
        content, usage = get_openai_completion(prompt, MODEL, json_output=json_output)
    try:
        _record_usage(MODEL, usage)
    except Exception:
        pass
    return content

# ───────────────────────── models ─────────────────────────

@dataclass
class AnswerLocator:
    type: str  # "paragraph_after" | "paragraph" | "table_cell"
    paragraph_index: Optional[int] = None   # for paragraph / paragraph_after
    offset: int = 0                         # how many paragraphs after the anchor
    table_index: Optional[int] = None       # for table_cell
    row: Optional[int] = None
    col: Optional[int] = None

@dataclass
class QASlot:
    id: str
    question_text: str
    answer_locator: AnswerLocator
    answer_type: str = "text"  # text | multiple_choice | file | table | checkbox | multi-select | date | number
    confidence: float = 0.5
    meta: Dict[str, Any] = None

# ───────────────────────── utils ─────────────────────────

def _iter_block_items(doc: docx.document.Document):
    """Yield paragraphs and tables in document order."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    parent = doc.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def _is_blank_para(p: Paragraph) -> bool:
    t = (p.text or "").strip()
    if t == "":
        return True
    # underscore lines or placeholders like [Insert response]
    if re.fullmatch(r"_+\s*", t):
        return True
    if re.fullmatch(r"\[(?:insert|enter|provide)[^\]]*\]", t.lower()):
        return True
    # check if any run is explicitly underlined without real text
    if any(r.text and r.underline for r in p.runs) and len(t.replace("_","").strip()) == 0:
        return True
    return False

QUESTION_PHRASES = (
    "please describe", "please provide", "explain", "detail", "outline",
    "how do you", "how will you", "what is your", "what are your", "do you",
    "can you", "does your", "have you", "who", "when", "where", "why", "which"
)

# ─────────────────── outline / context helpers ───────────────────
_ENUM_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"(?:\(?\d+(?:\.\d+)*\)?[.)]?)|"     # 1   1.1   2.3.4   (1)   1)
    r"(?:[A-Za-z][.)])|"                      # a)   A)   a.   A.
    r"(?:\([A-Za-z0-9]+\))"                 # (a)  (A)  (i)  (1)
    r")\s+"
)

def strip_enum_prefix(text: str) -> str:
    """Remove a single leading enumeration token like '1.1 ', '(a) ', 'A) '."""
    return _ENUM_PREFIX_RE.sub("", (text or ""), count=1)

def derive_outline_hint_and_level(text: str):
    """
    Return (hint, level) derived from visible enumeration.
    Examples: '1.1'→('1.1',2) ; '3)'→('3',1) ; '(a)'→('a',1)
    """
    t = (text or "").strip()
    m = re.match(r"^\s*(\d+(?:\.\d+)+)[.)]?\s+", t)
    if m:
        num = m.group(1)
        return num, num.count(".") + 1
    m = re.match(r"^\s*\(?(\d+)\)?[.)]?\s+", t)
    if m:
        return m.group(1), 1
    m = re.match(r"^\s*\(?([A-Za-z])\)?[.)]?\s+", t)
    if m:
        return m.group(1), 1
    return None, None

def paragraph_level_from_numbering(p: Paragraph):
    """Return Word automatic numbering level (1‑based) if present."""
    try:
        numPr = p._p.pPr.numPr
        if numPr is None:
            return None
        ilvl = numPr.ilvl.val if numPr.ilvl is not None else None
        if ilvl is None:
            return None
        return int(ilvl) + 1
    except Exception:
        return None

def heading_chain(blocks, upto_block, max_back=80):
    """Return list of Heading‑style texts (top→nearest) above a block index."""
    chain = []
    for bi in range(max(0, upto_block - max_back), upto_block):
        b = blocks[bi]
        if isinstance(b, Paragraph):
            try:
                style = (b.style.name or "").lower()
            except Exception:
                style = ""
            if style.startswith("heading"):
                txt = (b.text or "").strip()
                if txt:
                    chain.append(txt)
    return chain

def _looks_like_question(text: str) -> bool:
    t_raw = (text or "").strip()
    if not t_raw:
        return False

    # Presence of a question mark anywhere is a strong signal
    if "?" in t_raw:
        return True

    # Remove enumeration prefix (e.g. '1.1 ', '(a) ') before heuristic checks
    t = strip_enum_prefix(t_raw).strip()
    t_lower = t.lower()

    # Question/request cues at the start of the text
    if any(t_lower.startswith(phrase) for phrase in QUESTION_PHRASES):
        return True

    # Question/request cues appearing anywhere in the text
    if any(phrase in t_lower for phrase in QUESTION_PHRASES):
        return True

    if t_raw.lower().startswith(("question:", "prompt:", "rfp question:")):
        return True

    # If original had an enumeration token and contains a cue anywhere
    if _ENUM_PREFIX_RE.match(t_raw) and any(phrase in t_lower for phrase in QUESTION_PHRASES):
        return True

    return False

def _para_style_name(p: Paragraph) -> str:
    try:
        return p.style.name or ""
    except Exception:
        return ""

def _table_cell_text(table: Table, r: int, c: int) -> str:
    try:
        return (table.cell(r, c).text or "").strip()
    except Exception:
        return ""

# ─────────────────── rich extraction helpers (format-aware) ───────────────────

def _run_to_markup(r) -> str:
    t = (r.text or "")
    if not t:
        return ""
    # Wrap with simple tags to preserve formatting signal
    if r.underline:
        t = f"<u>{t}</u>"
    if r.italic:
        t = f"<i>{t}</i>"
    if r.bold:
        t = f"<b>{t}</b>"
    return t

def _paragraph_rich_text(p: Paragraph) -> str:
    if not p.runs:
        return (p.text or "")
    parts = []
    for r in p.runs:
        parts.append(_run_to_markup(r))
    return "".join(parts) or (p.text or "")

def _cell_rich_text(cell) -> str:
    # Join paragraphs inside a cell, preserving run markup
    lines = []
    for par in cell.paragraphs:
        lines.append(_paragraph_rich_text(par))
    return "\n".join(lines).strip()

def _para_alignment(p: Paragraph) -> str:
    try:
        al = p.alignment
        if al is None:
            return "NONE"
        return str(al)  # Enum repr is fine
    except Exception:
        return "NONE"

def _para_num_info(p: Paragraph) -> Tuple[Optional[int], Optional[int]]:
    """
    Returns (numId, ilvl) if numbering is present, else (None, None).
    """
    try:
        numPr = p._p.pPr.numPr  # type: ignore
        if numPr is None:
            return (None, None)
        numId = numPr.numId.val if numPr.numId is not None else None
        ilvl = numPr.ilvl.val if numPr.ilvl is not None else None
        return (numId, ilvl)
    except Exception:
        return (None, None)

def _para_indent_info(p: Paragraph) -> Tuple[int, int]:
    """Return left and first-line indent in points (0 if absent)."""
    try:
        pf = p.paragraph_format
        left = int(pf.left_indent.pt) if pf.left_indent else 0
        first = int(pf.first_line_indent.pt) if pf.first_line_indent else 0
        return left, first
    except Exception:
        return 0, 0

def _render_rich_excerpt(blocks: List[Union[Paragraph, Table]], start_index: int = 0) -> Tuple[str, Dict[int, int]]:
    """
    Produce a linearized, format-aware representation with global block indices.
    For tables, also return a map {global_block_index -> 0-based table_index}.
    """
    lines: List[str] = []
    table_idx_map: Dict[int, int] = {}
    running_table_index = 0
    for gi, b in enumerate(blocks, start=start_index):
        if isinstance(b, Paragraph):
            style = _para_style_name(b)
            align = _para_alignment(b)
            numId, ilvl = _para_num_info(b)
            left, first = _para_indent_info(b)
            raw = b.text or ""
            leading_ws = len(raw) - len(raw.lstrip(" \t"))
            rich = _paragraph_rich_text(b)
            lines.append(
                f"B[{gi}] PARAGRAPH style='{style}' align={align} numId={numId} ilvl={ilvl} "
                f"left={left} first={first} leading_ws={leading_ws}"
            )
            lines.append(f"B[{gi}] TEXT: {rich if rich else raw}")
        elif isinstance(b, Table):
            # Table header line
            try:
                rows = len(b.rows)
                cols = len(b.columns)
            except Exception:
                rows, cols = 0, 0
            table_idx_map[gi] = running_table_index
            lines.append(f"B[{gi}] TABLE rows={rows} cols={cols} (table_index={running_table_index})")
            for r in range(rows):
                row_line = []
                for c in range(cols):
                    try:
                        cell_text = _cell_rich_text(b.cell(r, c))
                    except Exception:
                        cell_text = ""
                    lines.append(f"B[{gi}] [{r},{c}] TEXT: {cell_text}")
            running_table_index += 1
    excerpt = "\n".join(lines)
    return excerpt, table_idx_map


def _render_structured_excerpt(blocks: List[Union[Paragraph, Table]], start_index: int = 0) -> str:
    """Return JSON string describing blocks for precise LLM inspection."""
    items: List[Dict[str, Any]] = []
    for gi, b in enumerate(blocks, start=start_index):
        if isinstance(b, Paragraph):
            style = _para_style_name(b)
            numId, ilvl = _para_num_info(b)
            left, first = _para_indent_info(b)
            items.append(
                {
                    "index": gi,
                    "type": "paragraph",
                    "text": b.text or "",
                    "style": style,
                    "numId": numId,
                    "ilvl": ilvl,
                    "left": left,
                    "first": first,
                }
            )
        elif isinstance(b, Table):
            try:
                rows = len(b.rows)
                cols = len(b.columns)
            except Exception:
                rows, cols = 0, 0
            cells: List[Dict[str, Any]] = []
            for r in range(rows):
                for c in range(cols):
                    try:
                        cell_text = _cell_rich_text(b.cell(r, c))
                    except Exception:
                        cell_text = ""
                    cells.append({"r": r, "c": c, "text": cell_text})
            items.append(
                {
                    "index": gi,
                    "type": "table",
                    "rows": rows,
                    "cols": cols,
                    "cells": cells,
                }
            )
    return json.dumps({"blocks": items}, ensure_ascii=False)

# ─────────────────── answer-type heuristics ───────────────────

_FILE_KWS = (
    "attach",
    "attachment",
    "upload",
    "enclose",
    "file",
    "document",
)
_TABLE_KWS = (
    "table",
    "spreadsheet",
    "complete the table",
    "fill in the table",
    "table below",
)
_MC_KWS = (
    "select",
    "choose",
    "check the box",
    "checkbox",
    "tick",
    "multiple choice",
    "which of the following",
)
_CHECKBOX_CHARS = "☐☑☒□■✓✔✗✘"


def llm_extract_mc_choices(blocks: List[Union[Paragraph, Table]], q_block: int) -> List[Dict[str, object]]:
    """Use an LLM to guess multiple-choice options when heuristics fail."""
    if FRAMEWORK != "openai" or not os.getenv("OPENAI_API_KEY"):
        dbg("llm_extract_mc_choices unavailable: framework!=openai or missing API key")
        return []

    question = ""
    if isinstance(blocks[q_block], Paragraph):
        question = blocks[q_block].text or ""
    dbg(f"llm_extract_mc_choices for q_block {q_block}: '{question}'")

    following: List[str] = []
    for nb in blocks[q_block + 1 : q_block + 10]:
        if isinstance(nb, Paragraph):
            following.append(nb.text or "")
        else:
            break
    context = "\n".join(following)
    dbg(f"Context lines after question: {len(following)}")

    template = read_prompt("mc_llm_scan")
    prompt = template.format(question=question, context=context)
    if SHOW_TEXT:
        print("\n--- PROMPT (llm_extract_mc_choices) ---")
        print(prompt)
        print("--- END PROMPT ---\n")
    try:
        resp = _call_llm(prompt, json_output=True)
        if SHOW_TEXT:
            print("\n--- COMPLETION (llm_extract_mc_choices) ---")
            print(resp)
            print("--- END COMPLETION ---\n")
        options = json.loads(resp)
        dbg(f"LLM suggested options: {options}")
    except Exception as e:
        dbg(f"llm_extract_mc_choices error: {e}")
        return []

    if not isinstance(options, list):
        dbg("LLM response was not a list of options")
        return []

    choices: List[Dict[str, object]] = []
    for opt in options:
        if not isinstance(opt, str):
            continue
        # try to locate paragraph containing this option text
        opt_low = opt.lower()
        for offset, nb in enumerate(blocks[q_block + 1 : q_block + 10], start=1):
            if isinstance(nb, Paragraph) and opt_low in (nb.text or "").lower():
                choices.append({"text": opt.strip(), "prefix": "", "block_index": q_block + offset})
                dbg(
                    f"Matched option '{opt.strip()}' to block {q_block + offset}"
                )
                break
    dbg(f"Final choices from LLM: {choices}")
    return choices


def extract_mc_choices(blocks: List[Union[Paragraph, Table]], q_block: int) -> List[Dict[str, object]]:
    """Collect multiple choice options appearing after the question block.

    Each choice is returned as a dict with:
      - text:       cleaned option text
      - prefix:     leading marker/prefix (checkbox, enumeration, etc.)
      - block_index:index of the paragraph containing the option
    """
    choices: List[Dict[str, object]] = []
    for offset, nb in enumerate(blocks[q_block + 1 : q_block + 10], start=1):
        if not isinstance(nb, Paragraph):
            break
        txt = (nb.text or "").strip()
        if not txt:
            continue
        prefix = ""
        cleaned = txt
        if any(ch in txt for ch in _CHECKBOX_CHARS):
            m = re.match(rf"^[{_CHECKBOX_CHARS}]\s*", txt)
            if m:
                prefix = m.group(0)
                cleaned = txt[m.end():].strip()
        elif re.match(r"^\(\s*\)\s*", txt):
            m = re.match(r"^\(\s*\)\s*", txt)
            prefix = m.group(0)
            cleaned = txt[m.end():].strip()
        elif re.match(r"^\[\s*\]\s*", txt):
            m = re.match(r"^\[\s*\]\s*", txt)
            prefix = m.group(0)
            cleaned = txt[m.end():].strip()
        elif re.match(_ENUM_PREFIX_RE, txt):
            m = _ENUM_PREFIX_RE.match(txt)
            if m:
                prefix = m.group(0)
                cleaned = txt[m.end():].strip()
        else:
            break
        choices.append({
            "text": cleaned,
            "prefix": prefix,
            "block_index": q_block + offset,
        })
    if not choices:
        dbg("Heuristic MC extraction found no choices; invoking LLM")
        choices = llm_extract_mc_choices(blocks, q_block)
        dbg(f"LLM returned choices: {choices}")
    return choices


def infer_answer_type(question_text: str, blocks: List[Union[Paragraph, Table]], q_block: int) -> str:
    """Guess the expected answer format for a question.

    The heuristic uses keywords in the question text and looks ahead a few
    blocks to inspect formatting cues such as checkboxes or tables.
    """

    t = (question_text or "").strip().lower()
    # Keyword-based checks first
    if any(kw in t for kw in _FILE_KWS):
        return "file"
    if any(kw in t for kw in _TABLE_KWS):
        return "table"
    if any(kw in t for kw in _MC_KWS):
        return "multiple_choice"

    # Look ahead at subsequent blocks for visual cues
    for nb in blocks[q_block + 1 : q_block + 6]:
        if isinstance(nb, Table):
            try:
                if len(nb.rows) > 1 and len(nb.columns) > 1:
                    return "table"
            except Exception:
                pass
        if isinstance(nb, Paragraph):
            txt = (nb.text or "").strip()
            low = txt.lower()
            if any(ch in txt for ch in _CHECKBOX_CHARS):
                return "multiple_choice"
            if re.search(r"\[[x ]\]|\([x ]\)", txt):
                return "multiple_choice"
            if re.match(_ENUM_PREFIX_RE, txt) and not _looks_like_question(txt):
                return "multiple_choice"
            if "yes" in low and "no" in low and len(low.split()) <= 4:
                return "multiple_choice"
    return "text"

# ─────────────────── rule-based detectors ───────────────────

def detect_para_question_with_blank(blocks: List[Union[Paragraph, Table]]) -> List[QASlot]:
    slots: List[QASlot] = []
    p_index = -1  # count paragraphs so we can locate
    for i, b in enumerate(blocks):
        if isinstance(b, Paragraph):
            p_index += 1
            text = (b.text or "").strip()
            if not _looks_like_question(text):
                continue

            # Look ahead up to 3 blocks for blank area or underscores or a 1x1 empty table
            conf_base = 0.6
            style = _para_style_name(b)
            if "Question" in style:
                conf_base += 0.1

            # 1) blank paragraphs
            for j in range(1, 4):
                if i + j >= len(blocks):
                    break
                nb = blocks[i + j]
                if isinstance(nb, Paragraph):
                    nb_text = (nb.text or "").strip()
                    if _looks_like_question(nb_text):
                        break
                    if _is_blank_para(nb):
                        lvl_num = paragraph_level_from_numbering(b)
                        hint, hint_level = derive_outline_hint_and_level(text)
                        ctx_level = lvl_num or hint_level
                        slots.append(QASlot(
                            id=f"slot_{uuid.uuid4().hex[:8]}",
                            question_text=text,
                            answer_locator=AnswerLocator(type="paragraph", paragraph_index=p_index + j),
                            answer_type=infer_answer_type(text, blocks, i),
                            confidence=min(0.95, conf_base + 0.2),
                            meta={
                                "detector": "para_blank_after",
                                "q_paragraph_index": p_index,
                                "q_block": i,
                                "q_style": style,
                                "outline": {"level": ctx_level, "hint": hint}
                            }
                        ))
                        break
                if isinstance(nb, Table):
                    try:
                        if len(nb.rows) == 1 and len(nb.columns) == 1:
                            cell_text = (nb.cell(0, 0).text or "").strip()
                            if cell_text == "":
                                t_idx = _running_table_index(blocks, i + j)
                                lvl_num = paragraph_level_from_numbering(b)
                                hint, hint_level = derive_outline_hint_and_level(text)
                                ctx_level = lvl_num or hint_level
                                slots.append(QASlot(
                                    id=f"slot_{uuid.uuid4().hex[:8]}",
                                    question_text=text,
                                    answer_locator=AnswerLocator(
                                        type="table_cell", table_index=t_idx, row=0, col=0
                                    ),
                                    answer_type=infer_answer_type(text, blocks, i),
                                    confidence=min(0.9, conf_base + 0.15),
                                    meta={
                                        "detector": "para_then_empty_1x1_table",
                                        "q_paragraph_index": p_index,
                                        "q_block": i,
                                        "outline": {"level": ctx_level, "hint": hint}
                                    }
                                ))
                                break
                    except Exception:
                        pass
    return slots

def _running_table_index(blocks: List[Union[Paragraph, Table]], upto: int) -> int:
    """Return table index counting from 0 in document order up to position `upto`."""
    t = 0
    for k, b in enumerate(blocks[:upto+1]):
        if isinstance(b, Table):
            if k == upto:
                return t
            t += 1
    return t

def detect_two_col_table_q_blank(blocks: List[Union[Paragraph, Table]]) -> List[QASlot]:
    slots: List[QASlot] = []
    table_counter = -1
    for i, b in enumerate(blocks):
        if isinstance(b, Table):
            table_counter += 1
            if len(b.columns) != 2:
                continue
            # Header detection
            header_left = _table_cell_text(b, 0, 0).lower()
            header_right = _table_cell_text(b, 0, 1).lower()
            has_header = any(k in header_left for k in ("question", "prompt")) or any(k in header_right for k in ("answer", "response"))
            start_row = 1 if has_header else 0

            for r in range(start_row, len(b.rows)):
                left = _table_cell_text(b, r, 0)
                right = _table_cell_text(b, r, 1)
                if not left:
                    continue
                # Question-like left, empty right
                if _looks_like_question(left) and right == "":
                    conf = 0.8 if has_header else 0.7
                    slots.append(QASlot(
                        id=f"slot_{uuid.uuid4().hex[:8]}",
                        question_text=left.strip(),
                        answer_locator=AnswerLocator(type="table_cell", table_index=table_counter, row=r, col=1),
                        answer_type=infer_answer_type(left, blocks, i),
                        confidence=conf,
                        meta={"detector": "table_2col_q_left_blank_right", "has_header": has_header}
                    ))
    return slots

def detect_response_label_then_blank(blocks: List[Union[Paragraph, Table]]) -> List[QASlot]:
    slots: List[QASlot] = []
    p_index = -1
    for i, b in enumerate(blocks):
        if isinstance(b, Paragraph):
            p_index += 1
            t = (b.text or "").strip()
            # Match patterns like "Response:" or "Answer:"
            if re.match(r"^(Response|Answer)\s*:\s*$", t, flags=re.IGNORECASE):
                # backtrack to find nearest prior question paragraph within 3 items
                q_text, q_idx = None, None
                back_p_idx = p_index
                for k in range(1, 4):
                    j = i - k
                    if j < 0: break
                    prev = blocks[j]
                    if isinstance(prev, Paragraph):
                        back_p_idx -= 1
                        if _looks_like_question((prev.text or "").strip()):
                            q_text = (prev.text or "").strip()
                            q_idx = back_p_idx
                            break
                # look forward to find first blank paragraph
                if q_text is not None:
                    for j in range(1, 4):
                        if i + j >= len(blocks): break
                        nb = blocks[i + j]
                        if isinstance(nb, Paragraph):
                            nb_text = (nb.text or "").strip()
                            if _looks_like_question(nb_text):
                                break
                            if _is_blank_para(nb):
                                lvl_num = paragraph_level_from_numbering(prev) if isinstance(prev, Paragraph) else None
                                hint, hint_level = derive_outline_hint_and_level(q_text)
                                ctx_level = lvl_num or hint_level
                                slots.append(QASlot(
                                    id=f"slot_{uuid.uuid4().hex[:8]}",
                                    question_text=q_text,
                                    answer_locator=AnswerLocator(type="paragraph", paragraph_index=p_index + j),
                                    answer_type=infer_answer_type(q_text, blocks, i - k),
                                    confidence=0.75,
                                    meta={
                                        "detector": "response_label_then_blank",
                                        "q_paragraph_index": q_idx,
                                        "q_block": (i - k),
                                        "outline": {"level": ctx_level, "hint": hint}
                                    }
                                ))
                                break
    return slots

# ─────────────────── optional LLM refiner ───────────────────

USE_LLM = True  # default is ON; can be disabled with --no-ai

def llm_refine(slots: List[QASlot], context_windows: List[str]) -> List[QASlot]:
    """
    Stub that could call an LLM to confirm/adjust low-confidence slots.
    Keep as no-op unless USE_LLM=True and you implement it.
    """
    if not USE_LLM:
        return slots

    refined: List[QASlot] = []
    for s, ctx in zip(slots, context_windows):
        if s.confidence >= 0.8:
            refined.append(s)
            continue
        template = read_prompt("docx_llm_refine")
        prompt = template.format(ctx=ctx)
        try:
            content = _call_llm(prompt, json_output=True)
            if SHOW_TEXT:
                print("\n--- PROMPT (llm_refine) ---")
                print(prompt)
                print("--- COMPLETION (llm_refine) ---")
                print(content)
                print("--- END COMPLETION ---\n")
            js = json.loads(content)
            if js.get("is_question"):
                s.confidence = max(s.confidence, 0.85)
        except Exception:
            pass
        refined.append(s)
    return refined

# ─────────────────── LLM paragraph‑scan fallback ───────────────────

def llm_scan_blocks(blocks: List[Union[Paragraph, Table]], model: str = MODEL) -> List[QASlot]:
    """If rule‑based detectors find nothing, let an LLM propose Q→A blanks."""
    excerpt, table_idx_map = _render_rich_excerpt(blocks)
    dbg(f"llm_scan_blocks (rich): {len(excerpt)} chars, model={model}")
    dbg(f"Sending prompt to LLM (first 400 chars): {excerpt[:400]}...")

    template = read_prompt("docx_llm_scan_blocks")
    prompt = template.format(doc=excerpt)
    if SHOW_TEXT:
        print("\n--- PROMPT (llm_scan_blocks) ---")
        print(prompt)
        print("--- END PROMPT ---\n")
    try:
        content = _call_llm(
            prompt,
            json_output=True,
        )
        js = json.loads(content)
        cand = js.get("slots", []) or []
        if SHOW_TEXT:
            print("\n--- COMPLETION (llm_scan_blocks) ---")
            print(content)
            print("--- END COMPLETION ---\n")
        dbg(f"LLM returned {len(cand)} slot candidates")
        dbg(f"LLM raw slot candidates: {cand}")
    except Exception as e:
        dbg(f"LLM error: {e}")
        return []

    results: List[QASlot] = []
    for it in cand:
        dbg(f"Processing candidate: {it}")
        try:
            kind = (it.get("kind") or "").strip()
            if kind == "paragraph_after":
                q_block = int(it["question"]["block"])
                offset = max(1, min(3, int(it["answer"]["offset"])))
                # derive question text if possible
                if 0 <= q_block < len(blocks) and isinstance(blocks[q_block], Paragraph):
                    q_text = (blocks[q_block].text or "").strip()
                else:
                    q_text = ""
                results.append(QASlot(
                    id=f"slot_{uuid.uuid4().hex[:8]}",
                    question_text=q_text,
                    answer_locator=AnswerLocator(type="paragraph_after", paragraph_index=q_block, offset=offset),
                    answer_type=infer_answer_type(q_text, blocks, q_block),
                    confidence=0.6,
                    meta={"detector": "llm_rich", "block": q_block, "offset": offset}
                ))
                dbg(f"Appended slot from candidate: {results[-1]}")
            elif kind == "table_cell":
                q_block = int(it["question"]["block"])
                qr = int(it["question"]["row"])
                qc = int(it["question"]["col"])
                ab = int(it["answer"]["block"])
                ar = int(it["answer"]["row"])
                ac = int(it["answer"]["col"])
                # table index mapping
                t_index = table_idx_map.get(q_block)
                if t_index is None:
                    # if the model referenced a table block incorrectly, skip
                    continue
                # derive question text if possible
                q_text = ""
                try:
                    tbl = blocks[q_block]
                    if isinstance(tbl, Table):
                        q_text = (tbl.cell(qr, qc).text or "").strip()
                except Exception:
                    pass
                results.append(QASlot(
                    id=f"slot_{uuid.uuid4().hex[:8]}",
                    question_text=q_text,
                    answer_locator=AnswerLocator(type="table_cell", table_index=t_index, row=ar, col=ac),
                    answer_type=infer_answer_type(q_text, blocks, q_block),
                    confidence=0.65,
                    meta={"detector": "llm_rich", "q_block": q_block, "answer_block": ab, "row": ar, "col": ac}
                ))
                dbg(f"Appended slot from candidate: {results[-1]}")
        except Exception as e:
            dbg(f"Parse candidate error: {e}")
            continue
    return results

# ─────────────────── 2‑stage LLM helpers ───────────────────

def llm_detect_questions(
    blocks: List[Union[Paragraph, Table]],
    model: str = MODEL,
    chunk_size: int = 10,
) -> List[int]:
    """Return global block indices that look like questions using only the LLM."""
    # Pre-filter: remove empty blocks or paragraphs with fewer than 4 words
    filtered: List[Union[Paragraph, Table]] = []
    for b in blocks:
        if isinstance(b, Paragraph):
            # skip paragraphs with less than 4 words
            if len((b.text or "").split()) < 4:
                continue
        filtered.append(b)
    blocks = filtered
    found: List[int] = []
    for start in range(0, len(blocks), chunk_size):
        end = min(len(blocks), start + chunk_size)
        excerpt = _render_structured_excerpt(blocks[start:end], start_index=start)
        detect_template = read_prompt("docx_detect_questions")
        prompt = detect_template.format(excerpt=excerpt)
        if SHOW_TEXT:
            print("\n--- PROMPT (detect_questions) ---\n" + prompt + "\n--- END PROMPT ---\n")
        try:
            content = _call_llm(prompt, json_output=True)
            if SHOW_TEXT:
                print("\n--- COMPLETION (detect_questions) ---\n" + content + "\n--- END COMPLETION ---\n")
            js = json.loads(content)
            questions = [int(i) for i in js.get("questions", [])]
            dbg(f"Model returned JSON (detect_questions) chunk {start}-{end}: {js}")
            # --- Begin debug print for each block ---
            flagged = set(questions)
            for rel_idx, b in enumerate(blocks[start:end]):
                gi = start + rel_idx
                if isinstance(b, Paragraph):
                    text = b.text or ""
                elif isinstance(b, Table):
                    # Combine all cell texts
                    cell_texts = []
                    try:
                        for row in b.rows:
                            for cell in row.cells:
                                cell_texts.append(_cell_rich_text(cell))
                        text = " | ".join(cell_texts)
                    except Exception:
                        text = ""
                else:
                    text = ""
                if gi in flagged:
                    dbg(f"Block {gi}: FLAGGED as question -> {text}")
                else:
                    dbg(f"Block {gi}: skipped -> {text}")
            # --- End debug print for each block ---
            found.extend(questions)
        except Exception as e:
            dbg(f"Error parsing detect_questions response (chunk {start}-{end}): {e}")

    unique_sorted = sorted(set(found))
    dbg(f"Questions indices extracted: {unique_sorted}")
    return unique_sorted


def _para_has_page_break(p: Paragraph) -> bool:
    """Return True if paragraph contains a hard page break."""
    for r in p.runs:
        try:
            for br in r._r.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    return True
        except Exception:
            continue
    return False


def _blocks_to_text_pages(blocks: List[Union[Paragraph, Table]]) -> List[str]:
    """Convert DOCX blocks into a list of page-level plain-text strings."""
    pages: List[str] = []
    current: List[str] = []
    for b in blocks:
        if isinstance(b, Paragraph):
            current.append(b.text or "")
            if _para_has_page_break(b):
                pages.append("\n".join(current).strip())
                current = []
        elif isinstance(b, Table):
            rows = []
            for row in b.rows:
                row_text = "\t".join((cell.text or "").strip() for cell in row.cells)
                rows.append(row_text)
            current.append("\n".join(rows))
    if current or not pages:
        pages.append("\n".join(current).strip())
    return pages


def _find_block_index_for_question(q: str, blocks: List[Union[Paragraph, Table]]) -> Optional[int]:
    """Find the global block index whose text contains the question string."""
    norm = re.sub(r"\s+", " ", q.lower().strip())
    for idx, b in enumerate(blocks):
        if isinstance(b, Paragraph):
            txt = re.sub(r"\s+", " ", (b.text or "").lower())
            if norm and norm in txt:
                return idx
        elif isinstance(b, Table):
            try:
                for row in b.rows:
                    for cell in row.cells:
                        txt = re.sub(r"\s+", " ", (cell.text or "").lower())
                        if norm and norm in txt:
                            return idx
            except Exception:
                continue
    return None


def llm_detect_questions_raw_text(
    blocks: List[Union[Paragraph, Table]],
    existing_questions: Set[str],
    model: str = MODEL,
    buffer: int = 200,
) -> List[int]:
    """Use page-level plain text to find additional question block indices."""
    pages = _blocks_to_text_pages(blocks)
    extra_blocks: List[int] = []
    for i, page in enumerate(pages):
        context = page
        if i > 0:
            context = pages[i - 1][-buffer:] + "\n" + context
        if i + 1 < len(pages):
            context = context + "\n" + pages[i + 1][:buffer]
        template = read_prompt("extract_questions")
        prompt = template.format(context=context)
        try:
            content = _call_llm(prompt)
            lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
        except Exception as e:
            dbg(f"Error in raw text detection page {i}: {e}")
            continue
        for q in lines:
            norm = q.lower().strip()
            if not norm or norm in existing_questions:
                continue
            idx = _find_block_index_for_question(q, blocks)
            if idx is not None:
                extra_blocks.append(idx)
                existing_questions.add(norm)
    return sorted(set(extra_blocks))


async def llm_locate_answer(blocks: List[Union[Paragraph, Table]], q_block: int, window: int = 3, model: str = MODEL) -> Optional[AnswerLocator]:
    """Given a question block index, ask the LLM to pick best answer location within ±window."""

    # Build context window
    start = max(0, q_block - window)
    end = min(len(blocks), q_block + window + 1)
    local_blocks = blocks[start:end]
    excerpt, table_idx_map = _render_rich_excerpt(local_blocks)
    template = read_prompt("docx_locate_answer")
    prompt = template.format(start=start, excerpt=excerpt)
    if SHOW_TEXT:
        print(f"\n--- PROMPT (locate_answer q_block={q_block}) ---\n" + prompt + "\n--- END PROMPT ---\n")
    try:
        content = await asyncio.to_thread(_call_llm, prompt, True)
    except Exception as e:
        dbg(f"LLM error (locate_answer q_block={q_block}): {e}")
        return None
    if SHOW_TEXT:
        print(
            f"\n--- COMPLETION (locate_answer q_block={q_block}) ---\n" + content + "\n--- END COMPLETION ---\n"
        )
    try:
        js = json.loads(content)
        dbg(f"Model returned JSON (locate_answer q_block={q_block}): {js}")
        kind = js.get("kind")
        if kind == "paragraph_after":
            offset = max(1, min(3, int(js.get("offset", 1))))
            locator = AnswerLocator(type="paragraph_after", paragraph_index=q_block, offset=offset)
            dbg(f"Mapped model response to AnswerLocator (paragraph_after): {locator}")
            return locator
        elif kind == "table_cell":
            row = int(js.get("row", 0))
            col = int(js.get("col", 0))
            # map local q_block (0) back to global q_block to get table index
            global_block_index = q_block
            _, table_idx_map_global = _render_rich_excerpt(blocks)
            t_index = table_idx_map_global.get(global_block_index)
            locator = AnswerLocator(type="table_cell", table_index=t_index, row=row, col=col)
            dbg(f"Mapped model response to AnswerLocator (table_cell): {locator}")
            return locator
    except Exception as e:
        dbg(f"Error parsing locate_answer for q_block {q_block}: {e}")
        return None


async def llm_assess_context(blocks: List[Union[Paragraph, Table]], q_block: int, model: str = MODEL) -> bool:
    """Return True if the question likely depends on previous context."""

    start = max(0, q_block - 2)
    end = min(len(blocks), q_block + 1)
    local_blocks = blocks[start:end]
    excerpt, _ = _render_rich_excerpt(local_blocks)
    template = read_prompt("docx_assess_context")
    prompt = template.format(local_index=q_block - start, excerpt=excerpt)
    try:
        content = await asyncio.to_thread(_call_llm, prompt, True)
    except Exception as e:
        dbg(f"LLM error (assess_context q_block={q_block}): {e}")
        return False
    try:
        js = json.loads(content)
        return bool(js.get("needs_context"))
    except Exception:
        return False

# ───────────────────────── pipeline ─────────────────────────

def extract_slots_from_docx(path: str) -> Dict[str, Any]:
    doc = docx.Document(path)
    blocks = list(_iter_block_items(doc))

    # Split any Paragraph with explicit line breaks into separate blocks
    expanded_blocks: List[Union[Paragraph, Table]] = []
    for b in blocks:
        if isinstance(b, Paragraph) and "\n" in (b.text or ""):
            for line in b.text.splitlines():
                # Clone the XML element and create a new Paragraph for each line
                p = Paragraph(b._p, doc)
                # Remove all existing runs from the clone
                for r in list(p.runs):
                    p._p.remove(r._r)
                # Add a single run containing just this line
                p.add_run(line)
                expanded_blocks.append(p)
        else:
            expanded_blocks.append(b)
    blocks = expanded_blocks

    dbg(f"extract_slots_from_docx: USE_LLM={USE_LLM}")
    dbg(f"Total blocks: {len(blocks)}")

    # If LLM mode (USE_LLM) is active, skip all rule-based detectors entirely.
    if USE_LLM:
        if FRAMEWORK == "openai" and not os.getenv("OPENAI_API_KEY"):
            raise RuntimeError("OPENAI_API_KEY not set; cannot run in pure AI mode.")
        if FRAMEWORK == "aladdin":
            required = [
                "aladdin_studio_api_key",
                "defaultWebServer",
                "aladdin_user",
                "aladdin_passwd",
            ]
            missing = [v for v in required if not os.getenv(v)]
            if missing:
                raise RuntimeError(
                    f"Missing environment variables for aladdin framework: {', '.join(missing)}"
                )
        llm_model = MODEL

        # Heuristic first: flag paragraphs containing '?' or the word 'please'
        q_indices_heur: List[int] = []
        for i, b in enumerate(blocks):
            if isinstance(b, Paragraph) and b.text:
                raw = b.text.strip()
                low = raw.lower()
                reason = None
                if '?' in raw:
                    reason = "contains '?'"
                elif 'please' in low:
                    reason = "contains 'please'"
                if reason:
                    dbg(f"Heuristic flagged block {i}: {reason} -> {raw}")
                    q_indices_heur.append(i)

        # Prepare blocks for LLM: those not already flagged
        remaining_blocks = []
        global_to_local = {}
        for gi, b in enumerate(blocks):
            if gi not in q_indices_heur:
                local_idx = len(remaining_blocks)
                remaining_blocks.append(b)
                global_to_local[local_idx] = gi

        # LLM detection on remaining blocks (chunk size default)
        local_q_indices = llm_detect_questions(remaining_blocks, model=llm_model)

        # Map local indices back to global
        q_indices_llm = [global_to_local[l] for l in local_q_indices if l in global_to_local]

        # Combine heuristic and LLM results
        q_indices = sorted(set(q_indices_heur + q_indices_llm))

        async def process_q_block(qb: int) -> Optional[QASlot]:
            try:
                loc = await llm_locate_answer(blocks, qb, window=3, model=llm_model)
                if loc is None:
                    return None
                q_text = (blocks[qb].text if isinstance(blocks[qb], Paragraph) else "").strip()
                slot_obj = QASlot(
                    id=f"slot_{uuid.uuid4().hex[:8]}",
                    question_text=q_text,
                    answer_locator=loc,
                    answer_type=infer_answer_type(q_text, blocks, qb),
                    confidence=0.6,
                    meta={"detector": "two_stage", "q_block": qb}
                )
                # Enrich slot_obj.meta with outline
                q_par = blocks[qb] if isinstance(blocks[qb], Paragraph) else None
                lvl_num = paragraph_level_from_numbering(q_par) if q_par else None
                hint, hint_level = derive_outline_hint_and_level(q_text)
                slot_obj.meta["outline"] = {"level": lvl_num or hint_level, "hint": hint}
                slot_obj.meta["needs_context"] = await llm_assess_context(blocks, qb, model=llm_model)
                dbg(f"Created QASlot from q_block {qb}: {asdict(slot_obj)}")
                return slot_obj
            except Exception as e:
                dbg(f"Error processing q_block {qb}: {e}")
                return None

        async def gather_slots() -> List[Optional[QASlot]]:
            tasks = [asyncio.create_task(process_q_block(qb)) for qb in q_indices]
            return await asyncio.gather(*tasks)

        slots = [s for s in asyncio.run(gather_slots()) if s]
        existing_qtexts = {s.question_text.strip().lower() for s in slots}
        extra_q_blocks = llm_detect_questions_raw_text(blocks, existing_qtexts, model=llm_model)
        for qb in extra_q_blocks:
            extra_slot = asyncio.run(process_q_block(qb))
            if extra_slot:
                if extra_slot.meta is None:
                    extra_slot.meta = {}
                extra_slot.meta["detector"] = "raw_text"
                slots.append(extra_slot)
        if not slots:  # fallback to legacy single scan
            slots = llm_scan_blocks(blocks, model=llm_model)
    else:
        # Fallback legacy rule‑based path (only when LLM explicitly disabled)
        slots: List[QASlot] = []
        for detector in (detect_para_question_with_blank,
                         detect_two_col_table_q_blank,
                         detect_response_label_then_blank):
            slots.extend(detector(blocks))
        # optional: refine if we still want refinement when LLM disabled (keep off)
    
    for s in slots:
        if s.answer_type == "multiple_choice":
            qb = (s.meta or {}).get("q_block")
            if qb is not None:
                choices = extract_mc_choices(blocks, qb)
                if choices:
                    if s.meta is None:
                        s.meta = {}
                    s.meta["choices"] = choices

    attach_context(slots, blocks)
    dbg(f"Slot count: {len(slots)}")
    dbg(f"Final payload preview: {json.dumps({'doc_type': 'docx', 'file': os.path.basename(path), 'slots': [asdict(s) for s in dedupe_slots(slots)]}, indent=2)[:1000]}")
    payload = {
        "doc_type": "docx",
        "file": os.path.basename(path),
        "slots": [asdict(s) for s in dedupe_slots(slots)]
    }#
    return payload

# ─────────────────── context attachment ───────────────────
def attach_context(slots: List[QASlot], blocks):
    """
    Populate slot.meta['context'] with: level, heading_chain and optional parent_*.
    """
    ordered = sorted(slots, key=lambda s: (s.meta or {}).get("q_block", 0))
    last_at_level = {}
    for s in ordered:
        qb = (s.meta or {}).get("q_block", 0)
        level = (s.meta or {}).get("outline", {}).get("level") or 1
        heads = heading_chain(blocks, qb)

        parent = None
        for l in range(level - 1, 0, -1):
            if l in last_at_level:
                parent = last_at_level[l]
                break

        ctx = {"level": int(level), "heading_chain": heads}
        if parent:
            ctx["parent_slot_id"] = parent.id
            ctx["parent_question_text"] = parent.question_text

        if s.meta is None:
            s.meta = {}
        s.meta["context"] = ctx
        last_at_level[level] = s

def dedupe_slots(slots: List[QASlot]) -> List[QASlot]:
    """Collapse slots with identical questions and overlapping locator ranges."""

    def norm_q(text: str) -> str:
        return strip_enum_prefix(text or "").strip().lower()

    def loc_range(slot: QASlot):
        loc = slot.answer_locator
        if loc.type == "table_cell":
            return ("cell", loc.table_index, loc.row, loc.col)
        if loc.type == "paragraph":
            return (loc.paragraph_index, loc.paragraph_index)
        if loc.type == "paragraph_after":
            start = (loc.paragraph_index or 0) + 1
            end = (loc.paragraph_index or 0) + loc.offset
            return (start, end)
        return None

    def more_specific(a: QASlot, b: QASlot) -> QASlot:
        ra, rb = loc_range(a), loc_range(b)
        if ra is None or rb is None:
            return a
        if len(ra) == 4 and len(rb) == 4:
            return a  # same cell → keep first
        la = ra[1] - ra[0]
        lb = rb[1] - rb[0]
        if la != lb:
            return a if la < lb else b
        pr = {"paragraph": 2, "paragraph_after": 1}
        return a if pr.get(a.answer_locator.type, 0) >= pr.get(b.answer_locator.type, 0) else b

    out: List[QASlot] = []
    for s in slots:
        nq = norm_q(s.question_text)
        r = loc_range(s)
        dup_idx = None
        for i, ex in enumerate(out):
            if norm_q(ex.question_text) != nq:
                continue
            er = loc_range(ex)
            if r is None or er is None:
                continue
            if len(r) == 4 and len(er) == 4:
                if r == er:
                    dup_idx = i
                    break
            elif len(r) == 2 and len(er) == 2:
                if not (r[1] < er[0] or r[0] > er[1]):
                    dup_idx = i
                    break
        if dup_idx is not None:
            out[dup_idx] = more_specific(out[dup_idx], s)
        else:
            out.append(s)
    return out

# ───────────────────────── CLI ─────────────────────────

def main():
    ap = argparse.ArgumentParser(description="Detect QA slots in a DOCX RFP.")
    ap.add_argument("docx_path", help="Path to .docx")
    ap.add_argument("-o", "--out", default=None, help="Write JSON to this path")
    ap.add_argument("--ai", action="store_true", help="(deprecated) AI mode is always on; flag kept for backward compatibility")
    ap.add_argument(
        "--debug",
        dest="debug",
        action="store_true",
        default=True,
        help="Print verbose debug info (default on)",
    )
    ap.add_argument(
        "--no-debug",
        dest="debug",
        action="store_false",
        help="Disable debug info",
    )
    ap.add_argument("--show-text", action="store_true", help="Dump full prompt and completion text for each API call (verbose)")
    ap.add_argument(
        "--framework",
        choices=["openai", "aladdin"],
        default=os.getenv("ANSWER_FRAMEWORK", "openai"),
        help="Which completion framework to use",
    )
    ap.add_argument(
        "--model",
        default=os.getenv("OPENAI_MODEL", "gpt-5-nano"),
        help="Model name for the chosen framework",
    )
    if len(sys.argv) == 1:
        ap.print_help()
        sys.exit(1)
    args = ap.parse_args()

    # Set debug global
    global DEBUG
    DEBUG = args.debug
    if DEBUG:
        print("### DEBUG MODE ON ###")
        print(f"[slot_finder] processing {args.docx_path}")

    global SHOW_TEXT
    SHOW_TEXT = args.show_text

    # AI is always enabled; --ai is legacy, --no-ai removed.
    global USE_LLM
    USE_LLM = True  # AI is always on; --ai is optional/legacy

    global FRAMEWORK, MODEL
    FRAMEWORK = args.framework
    MODEL = args.model

    if FRAMEWORK == "openai" and not os.getenv("OPENAI_API_KEY"):
        print("Error: OPENAI_API_KEY is not set (required for openai framework).", file=sys.stderr)
        sys.exit(1)
    if FRAMEWORK == "aladdin":
        required = ["aladdin_studio_api_key", "defaultWebServer", "aladdin_user", "aladdin_passwd"]
        missing = [v for v in required if not os.getenv(v)]
        if missing:
            print(
                f"Error: Missing environment variables for aladdin framework: {', '.join(missing)}",
                file=sys.stderr,
            )
            sys.exit(1)

    # Validate file existence
    if not os.path.isfile(args.docx_path):
        print(f"Error: File '{args.docx_path}' does not exist.", file=sys.stderr)
        sys.exit(1)
    # Validate .docx extension
    if not args.docx_path.lower().endswith(".docx"):
        print(f"Error: File '{args.docx_path}' does not have a .docx extension.", file=sys.stderr)
        sys.exit(1)
    try:
        if DEBUG:
            print("[slot_finder] extracting slots from DOCX")
        result = extract_slots_from_docx(args.docx_path)
        if DEBUG:
            print(f"[slot_finder] found {len(result.get('slots', []))} slots")
    except Exception as e:
        print(f"Error: Failed to process DOCX file '{args.docx_path}'. The file may be invalid or corrupted.\nDetails: {e}", file=sys.stderr)
        sys.exit(1)
    # Print token/cost summary if debug enabled
    if DEBUG:
        print("--- TOKEN / COST SUMMARY ---")
        print(f"Prompt tokens:  {TOTAL_INPUT_TOKENS}")
        print(f"Completion tokens: {TOTAL_OUTPUT_TOKENS}")
        print(f"Total estimated cost: ${TOTAL_COST_USD:.4f}")
    js = json.dumps(result, indent=2, ensure_ascii=False)
    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(js)
        if DEBUG:
            print(f"[slot_finder] wrote output to {args.out}")
        else:
            print(f"Wrote {args.out}")
    else:
        print(js)

if __name__ == "__main__":
    main()
