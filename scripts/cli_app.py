#!/usr/bin/env python3

"""
cli_app.py - RFP Responder CLI

Modes by input type:
- XLSX/XLS: detect questions from a worksheet template and write answers back to a new workbook
- DOCX (default): detect question "slots" in the template and apply answers into the same layout
- PDF/TXT (and DOCX with --doc-as-text): extract questions as text, answer, and build a Q/A Word report

Dependencies in this repo:
- llm.completions_client.CompletionsClient
- retrieval.vector_search.search
- documents.docx.slot_finder.extract_slots_from_docx
- documents.docx.apply_answers.apply_answers_to_docx
"""

from __future__ import annotations

import sys
import io
import os
import re
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Tuple, Dict, Any, Callable

import PyPDF2
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from openpyxl import load_workbook

# Ensure repository root is on sys.path
ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from backend.retrieval.stacks.faiss.structured_extraction.interpreter_sheet import collect_non_empty_cells

from backend.llm.completions_client import CompletionsClient
from backend.retrieval.vector_search import search
from backend.prompts import read_prompt

# DOCX slot finder + applier
from backend.documents.docx.slot_finder import extract_slots_from_docx
from backend.documents.docx.apply_answers import (
    apply_answers_to_docx,
    prepare_slots_with_answers,
)
from backend.documents.xlsx.apply_answers import write_excel_answers
from backend.documents.xlsx.slot_finder import ask_sheet_schema
from backend.answering.qa_engine import answer_question


PROMPTS = {
    "extract_questions": read_prompt("extract_questions"),
    "answer_search_context": read_prompt("answer_search_context"),
    "answer_llm": read_prompt("answer_llm_template"),
}

PRESET_INSTRUCTIONS = {
    "short": "Answer briefly in 1-2 sentences.",
    "medium": "Answer in one concise paragraph.",
    "long": "Answer in detail (up to one page).",
    "auto": "Answer using only the provided sources and choose an appropriate length.",
}


# Input text helpers (unchanged)
def load_input_text(path: Optional[str]) -> str:
    if path is None:
        print("[DEBUG] Reading text from stdin")
        return sys.stdin.read()
    p = Path(path)
    if not p.exists():
        print(f"[ERROR] File not found: {p}", file=sys.stderr)
        sys.exit(1)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        out = []
        with p.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                out.append(page.extract_text() or "")
        return "\n".join(out)
    if suffix in ".docx .doc":
        doc = Document(p)
        return "\n".join(par.text for par in doc.paragraphs)
    return p.read_text(encoding="utf-8")


# _______________________________________________________
# LLM question extraction (unchanged)
# _______________________________________________________


def extract_questions(text: str, llm: CompletionsClient) -> List[str]:
    tpl = PROMPTS["extract_questions"]
    prompt = tpl.format(text=text)
    print("[DEBUG] Extracting questions via LLM")
    resp = llm.get_completion(prompt).strip()
    lines = resp.splitlines()
    out: List[str] = []
    for line in lines:
        m = re.match(r"^\s*\d+\)\s+(.*)\s*$", line)
        if m:
            out.append(m.group(1).strip())
    print(f"[DEBUG] Extracted {len(out)} questions")
    return out


# Q/A DOCX report builder (unchanged)


def build_docx(
    questions: List[str],
    answers: List[str],
    comments: List[List[Tuple[str, str, str, float, str]]],
    include_comments: bool,
) -> bytes:
    print("[DEBUG] Building DOCX")
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    def _ensure_comments_part(document: Document):
        try:
            return document.part._comments_part
        except AttributeError:
            return document.part._add_comments_part()

    def _make_r(text: str, bold: bool = False):
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            b.set(Document().part.element.xmlns["w"] + "val", "1")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        return r

    com_part = None
    cid = 0

    def add_comment(run, lbl, src, snippet, score, date_str):
        nonlocal com_part, cid
        if com_part is None:
            com_part = _ensure_comments_part(doc)
        c = OxmlElement("w:comment")
        c.set(Document().part.element.xmlns["w"] + "id", str(cid))
        c.set(Document().part.element.xmlns["w"] + "author", "RFPBot")
        c.set(
            Document().part.element.xmlns["w"] + "date",
            datetime.utcnow().isoformat() + "Z",
        )
        for label, val, is_b in [
            ("Citation: ", f"[{lbl}]", True),
            ("Source: ", src, True),
            ("Date: ", date_str, True),
            ("Score: ", f"{score:.3f}", True),
            ("Quote: ", snippet, False),
        ]:
            p = OxmlElement("w:p")
            p.append(_make_r(label, bold=True))
            p.append(_make_r(val, bold=is_b))
            com_part.element.append(p)

        # anchor
        start = OxmlElement("w:commentRangeStart")
        start.set(Document().part.element.xmlns["w"] + "id", str(cid))
        end = OxmlElement("w:commentRangeEnd")
        end.set(Document().part.element.xmlns["w"] + "id", str(cid))
        ref = OxmlElement("w:commentReference")
        ref.set(Document().part.element.xmlns["w"] + "id", str(cid))

        parent = run._r.getparent()
        idx = parent.index(run._r)
        parent.insert(idx, start)
        parent.insert(idx + 1, end)
        parent.insert(idx + 2, ref)
        cid += 1

    for q, ans, cmts in zip(questions, answers, comments):
        pq = doc.add_paragraph()
        rq = pq.add_run("Q: ")
        rq.bold = True
        rq = pq.add_run(q)

        pa = doc.add_paragraph()
        ra = pa.add_run("A: ")
        ra.bold = True

        if include_comments:
            parts = re.split(r"(\[\d+\])", ans)
            for seg in parts:
                m = re.match(r"\[(\d+)\]$", seg)
                if m:
                    idx = int(m.group(1)) - 1
                    if idx < len(cmts):
                        lbl, src, snippet, score, date_str = cmts[idx]
                        run = pa.add_run(seg)
                        add_comment(run, lbl, src, snippet, score, date_str)
                    else:
                        pa.add_run(re.sub(r"\[\d+\]", "", seg))
                else:
                    pa.add_run(re.sub(r"\[\d+\]", "", seg))
        else:
            pa.add_run(re.sub(r"\[\d+\]", "", ans))

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


#
# Helpers for DOCX apply flow
#


def _make_docx_generator(
    *,
    search_mode: str,
    fund: Optional[str],
    k: int,
    length: Optional[str],
    approx_words: Optional[int],
    min_confidence: float,
    include_citations_in_text: bool,
    llm: CompletionsClient,
    extra_docs: Optional[List[str]] = None,
) -> Callable[[str], str]:
    """
    Returns a callable(question)->answer_text suitable for rfp_docx_apply_answers.apply_answers_tb_docx.
    """

    def gen(question: str) -> str:
        ans, cmts = answer_question(
            question,
            search_mode,
            fund,
            k,
            length,
            approx_words,
            min_confidence,
            llm,
            extra_docs=extra_docs,
        )
        if not include_citations_in_text:
            ans = re.sub(r"[\[\d+\]]", "", ans)
        return ans

    return gen


#
# Main CLI
#
def main():
    # Save current sys.path
    _original_sys_path = sys.path.copy()
    # Expand ~ to the home directory
    rfp_dir = os.path.expanduser("~/rfp_utils/rfp_utils")
    sys.path.insert(0, rfp_dir)
    # DOCX slot finder + applier
    from my_module import gen_answer as gen  # <-- lives in ~/rfp_utils/rfp_utils

    # Restore sys.path
    sys.path = _original_sys_path

    parser = argparse.ArgumentParser(description="RFP Responder CLI")
    parser.add_argument("input_file", help="PDF/DOCX/TXT/XLSX/XLS")

    parser.add_argument("--fund", required=False, help="Fund tag filter")
    group = parser.add_mutually_exclusive_group()
    group.add_argument(
        "--length", choices=["short", "medium", "long"], help="Preset length"
    )
    group.add_argument("--approx_words", type=int, help="Approximate word count")
    parser.add_argument("--no_comments", action="store_true", help="Disable citations")
    parser.add_argument(
        "--extra-doc",
        dest="extra_docs",
        action="append",
        help="Additional document to include via LLM search",
    )

    parser.add_argument(
        "--search_mode",
        choices=["answer", "question", "blend", "dual", "both"],
        default="dual",
    )
    parser.add_argument(
        "--llm_model", choices=["gpt-3.5-turbo", "gpt-4"], default="gpt-4o"
    )
    parser.add_argument(
        "-k", "--k_max_hits", type=int, default=6, help="Hits per question"
    )
    parser.add_argument(
        "--min_confidence", type=float, default=0.0, help="Min confidence threshold"
    )
    parser.add_argument("-o", "--output", help="Output file path")

    # New DOCX-specific flags
    parser.add_argument(
        "--docx-as-text",
        action="store_true",
        help="Treat docx like free text (build separate Q/A report) instead of applying in place.",
    )
    parser.add_argument(
        "--docx-write-mode",
        choices=["fill", "replace", "append"],
        default="fill",
        help="Write mode for DOCX apply flow",
    )
    parser.add_argument(
        "--slots", help="If provided (DOCX), also save detected slots JSON here"
    )

    args = parser.parse_args()

    # Prefer the Utilities-owned generator; fall back to the local one used for DOCX
    try:
        from my_module import gen_answer as _ext_gen  # provided by rfp_utils

        gen = _ext_gen
        gen_name = "rfp_utils.my_module:gen_answer"
        print(f"[DEBUG] Using external generator: {gen_name}")
    except Exception as e:
        print(f"[WARN] Falling back to built-in generator: {e}", file=sys.stderr)
        gen = _make_docx_generator(
            search_mode=args.search_mode,
            fund=args.fund,
            k=args.k_max_hits,
            length=args.length,
            approx_words=args.approx_words,
            min_confidence=args.min_confidence,
            include_citations_in_text=not args.no_comments,
            llm=CompletionsClient(model=os.environ.get("OPENAI_MODEL", "gpt-4o")),
            extra_docs=args.extra_docs,
        )
        gen_name = "cli_app:gen"

    infile = Path(args.input_file)
    if not infile.exists():
        print(f"[ERROR] Input not found: {infile}", file=sys.stderr)
        sys.exit(1)

    suffix = infile.suffix.lower()
    llm = CompletionsClient(model=args.llm_model)

    # --- Excel flow
    if suffix in (".xlsx", ".xls"):
        print(f"[DEBUG] Excel input detected: {infile}")
        cells = collect_non_empty_cells(infile)
        print(f"[DEBUG] Found {len(cells)} non-empty cells")

        # Debug: Print *all* cell contents - FIXED VERSION
        if isinstance(cells[0], dict):
            for i, cell in enumerate(cells[:10]):  # Show first 10 cells
                print(f"[DEBUG] Cell {i}: {cell}")
        else:
            for i, cell in enumerate(cells[:10]):  # Show first 10 cells
                print(
                    f"[DEBUG] Cell {i}: {getattr(cell, 'row', 'N/A')}, {getattr(cell, 'column', 'N/A')} = '{getattr(cell, 'value', 'N/A')}'"
                )

        if not cells:
            print("[ERROR] No non-empty cells found", file=sys.stderr)
            sys.exit(1)

        schema = ask_sheet_schema(infile)
        print(f"[DEBUG] Schema entries: {len(schema)}")

        # Debug: Print all schema entries
        for i, entry in enumerate(schema):
            print(f"[DEBUG] Schema Entry {i+1}:")
            if isinstance(entry, dict):
                for key, value in entry.items():
                    print(f"  {key}: {repr(value)}")
            else:
                print(f"  Entry type: {type(entry)}")
                print(f"  Entry content: {entry}")

        # Debug: Check what's being processed
        if len(schema) == 1 and schema[0] != {}:
            print("[DEBUG] Single schema entry detected, examining details:")
            for key, value in schema[0].items():
                print(f"[DEBUG] Schema key '{key}': {repr(value)}")

        answers: List[object] = []
        for i, entry in enumerate(schema, start=1):
            qtext = (entry.get("question_text") or "").strip()
            print(f"[DEBUG] Answering question {i}/{len(schema)} via {gen_name}")
            print(f"[DEBUG] Question text: '{qtext}'")
            ans = gen(qtext)
            answers.append(ans)

        out_path = (
            Path(args.output)
            if args.output
            else infile.with_name(infile.stem + "_answered.xlsx")
        )
        write_excel_answers(schema, answers, infile, out_path)
        print(f"[DEBUG] Wrote filled Excel to {out_path}")
        sys.exit(0)

    # — DOCX flow (apply answers into the template) ————————
    if suffix == ".docx" and not args.docx_as_text:
        print(f"[DEBUG] DOCX template detected: {infile}")
        # 1) Detect slots
        slots_payload = extract_slots_from_docx(str(infile))
        if args.slots:
            Path(args.slots).write_text(json.dumps(slots_payload), encoding="utf-8")

        # 2) Prefer external generator from Utilities (keeps CLI thin)
        try:
            gen_name = "rfp_utils.my_module:gen_answer"
            print("[DEBUG] Using external generator:", gen_name)
        except Exception as e:
            print(f"[WARN] Falling back to built-in generator ({e})", file=sys.stderr)
            gen = _make_docx_generator(
                search_mode=args.search_mode,
                fund=args.fund,
                k=args.max_hits,
                length=args.length,
                approx_words=args.approx_words,
                min_confidence=args.min_confidence,
                include_citations_in_text=not args.no_comments,
                llm=llm,
                extra_docs=args.extra_docs,
            )
            gen_name = "cli_app:rag_gen"

        # 3) Apply into a new docx
        out_path = (
            Path(args.output)
            if args.output
            else infile.with_name(infile.stem + "_answered.docx")
        )
        slots_list, generated = prepare_slots_with_answers(
            slots_payload,
            generator=gen,
            gen_name=gen_name,
        )
        _, summary = apply_answers_to_docx(
            docx_source=str(infile),
            slots=slots_list,
            mode=args.docx_write_mode,
            output_path=str(out_path),
        )
        summary["generated"] = generated
        print(f"Wrote {out_path}")
        print("[DEBUG] Apply summary:", summary)
        sys.exit(0)

    # ---- Text flow (PDF/TXT or DOCX with --docx-as-text)
    raw = load_input_text(str(infile))
    if not raw.strip():
        print("[ERROR] No text extracted", file=sys.stderr)
        sys.exit(1)

    questions = extract_questions(raw, llm)
    if not questions:
        print("[ERROR] No questions found", file=sys.stderr)
        sys.exit(1)

    answers: List[str] = []
    comments: List[List[Tuple[str, str, str, float, str]]] = []
    for i, q in enumerate(questions, start=1):
        print(f"[DEBUG] Answering question {i}/{len(questions)}")
        ans, cmts = answer_question(
            q,
            args.search_mode,
            args.fund,
            args.max_hits,
            args.length,
            args.approx_words,
            args.min_confidence,
            llm,
        )
        answers.append(ans)
        comments.append(cmts)

    qa_doc = build_docx(
        questions, answers, comments, include_comments=not args.no_comments
    )
    out_path = (
        Path(args.output)
        if args.output
        else infile.with_name(infile.stem + "_answered.docx")
    )
    out_path.write_bytes(qa_doc)
    print(f"[DEBUG] Wrote Q/A Word report to {out_path}")


# Small helper so we can safely write JSON even if ensure_ascii=False is wanted
def json_dump(obj: Any) -> str:
    import json

    return json.dumps(obj, indent=2, ensure_ascii=False)


if __name__ == "__main__":
    main()
