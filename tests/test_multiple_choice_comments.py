import pathlib
import sys

import pytest

pytest.importorskip("docx")
try:
    import docx.table  # noqa: F401
    import docx.text.paragraph  # noqa: F401
except ModuleNotFoundError:
    pytest.skip(
        "python-docx with table/text support is required for multiple-choice tests",
        allow_module_level=True,
    )

import docx

sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

from backend.documents.docx.apply_answers import iter_block_items, mark_multiple_choice
from backend.documents.docx.comments import ensure_comments_part


def test_mark_multiple_choice_adds_comment(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("() Yes")
    doc.add_paragraph("() No")
    path = tmp_path / "mc.docx"
    doc.save(path)

    doc = docx.Document(path)
    blocks = list(iter_block_items(doc))
    choices_meta = [
        {"block_index": 0, "text": "Yes", "prefix": "()"},
        {"block_index": 1, "text": "No", "prefix": "()"},
    ]
    mark_multiple_choice(doc, blocks, choices_meta, 0, comment_text="Evidence snippet")
    doc.save(path)

    reopened = docx.Document(path)
    part = ensure_comments_part(reopened)
    xml = part._element.xml
    assert "Evidence snippet" in xml
    assert "Source Text" in xml
    assert "<w:b/>" in xml
