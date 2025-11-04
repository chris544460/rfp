import os
import pathlib
import sys
from dataclasses import asdict

import pytest

sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

pytest.importorskip("docx")
try:
    import docx.table  # noqa: F401
    import docx.text.paragraph  # noqa: F401
except ModuleNotFoundError:
    pytest.skip(
        "python-docx with table/text support is required for workflow tests",
        allow_module_level=True,
    )

import docx
from backend.documents.docx.slot_finder import (
    detect_para_question_with_blank,
    detect_two_col_table_q_blank,
    detect_response_label_then_blank,
    _iter_block_items,
    attach_context,
    dedupe_slots,
    QASlot,
    AnswerLocator
)
from backend.documents.docx.apply_answers import apply_answers_to_docx

def build_slots(doc_path):
    doc = docx.Document(doc_path)
    blocks = list(_iter_block_items(doc))
    slots = []
    for detector in (
        detect_para_question_with_blank,
        detect_two_col_table_q_blank,
        detect_response_label_then_blank,
    ):
        slots.extend(detector(blocks))
    attach_context(slots, blocks)
    slots = dedupe_slots(slots)
    return {"doc_type": "docx", "file": os.path.basename(doc_path), "slots": [asdict(s) for s in slots]}

def test_pipeline_idempotent(tmp_path):
    answer_text = "Our approach is outstanding."
    doc = docx.Document()
    doc.add_paragraph("Please describe your approach.")
    doc.add_paragraph("")
    src = tmp_path / "orig.docx"
    doc.save(src)
    slots_payload = build_slots(src)
    slot_id = slots_payload["slots"][0]["id"]
    slots_with_answers = []
    for slot in slots_payload["slots"]:
        slot_copy = dict(slot)
        if slot_copy.get("id") == slot_id:
            slot_copy["answer"] = answer_text
        slots_with_answers.append(slot_copy)

    first = tmp_path / "out1.docx"
    apply_answers_to_docx(str(src), slots_with_answers, output_path=str(first))
    second = tmp_path / "out2.docx"
    apply_answers_to_docx(str(first), slots_with_answers, output_path=str(second))
    final_doc = docx.Document(second)
    texts = [p.text.strip() for p in final_doc.paragraphs]
    assert texts.count(answer_text) == 1

def test_dedupe_overlapping_ranges():
    s1 = QASlot(
        id="1",
        question_text="1. Please describe your approach.",
        answer_locator=AnswerLocator(type="paragraph", paragraph_index=5),
    )
    s2 = QASlot(
        id="2",
        question_text="Please describe your approach.",
        answer_locator=AnswerLocator(type="paragraph_after", paragraph_index=4, offset=2),
    )
    deduped = dedupe_slots([s1, s2])
    assert len(deduped) == 1


def test_blank_search_stops_before_next_question(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Question one?")
    doc.add_paragraph("Question two?")
    doc.add_paragraph("")
    src = tmp_path / "two_questions.docx"
    doc.save(src)
    loaded = docx.Document(src)
    blocks = list(_iter_block_items(loaded))
    slots = detect_para_question_with_blank(blocks)
    assert len(slots) == 1
    assert slots[0].question_text.strip() == "Question two?"
