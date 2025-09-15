import sys, pathlib

import pytest
from docx import Document

sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

from llm_doc_search import _extract_text_from_doc, search_uploaded_docs


class DummyLLM:
    def get_completion(self, prompt: str):
        if "Project Apollo is top secret." in prompt:
            return "YES: Project Apollo is top secret."
        if "Value 2" in prompt:
            return "YES: Value 2"
        return "NO"


def test_search_uploaded_docs_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Project Apollo is top secret.")
    path = tmp_path / "sample.docx"
    doc.save(path)

    llm = DummyLLM()
    hits = search_uploaded_docs("What is Project Apollo?", [str(path)], llm)
    assert hits
    assert "Project Apollo is top secret." in hits[0]["text"]


def test_search_uploaded_docs_docx_table(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Code"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Metric"
    table.cell(1, 1).text = "Value 2"
    path = tmp_path / "table.docx"
    doc.save(path)

    llm = DummyLLM()
    hits = search_uploaded_docs("What is the metric value?", [str(path)], llm)
    assert hits
    assert "Value 2" in hits[0]["text"]


def test_extract_docx_tables(tmp_path):
    doc = Document()
    doc.add_paragraph("Overview")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Row 1"
    table.cell(1, 1).text = "Value 2"
    path = tmp_path / "table.docx"
    doc.save(path)

    text = _extract_text_from_doc(str(path))
    assert "Overview" in text
    assert "[Table 1]" in text
    assert "| Header A | Header B |" in text
    assert "| Row 1 | Value 2 |" in text


def _build_pdf_with_table(path):
    pytest.importorskip("pdfplumber")
    pytest.importorskip("reportlab")

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("Table Overview", styles["Normal"]), Spacer(1, 12)]
    data = [["Header A", "Header B"], ["Row 1", "Value 2"]]
    tbl = Table(data)
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
    ]))
    story.append(tbl)
    doc.build(story)


def test_extract_pdf_tables(tmp_path):
    pytest.importorskip("pdfplumber")
    pytest.importorskip("reportlab")

    path = tmp_path / "table.pdf"
    _build_pdf_with_table(path)

    text = _extract_text_from_doc(str(path))
    assert "[Page 1]" in text
    assert "[Table 1 | Page 1]" in text
    assert "| Header A | Header B |" in text
    assert "| Row 1 | Value 2 |" in text
