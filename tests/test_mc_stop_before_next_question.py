import pathlib
import sys

import pytest

pytest.importorskip("docx")
try:
    import docx.table  # noqa: F401
    import docx.text.paragraph  # noqa: F401
except ModuleNotFoundError:
    pytest.skip(
        "python-docx with table/text support is required for docx slot finder tests",
        allow_module_level=True,
    )

import docx

# Ensure project root on path
sys.path.append(str(pathlib.Path(__file__).resolve().parents[1]))

from backend.documents.docx.slot_finder import extract_mc_choices, _iter_block_items


def test_extract_mc_choices_stops_before_next_question():
    doc = docx.Document()
    doc.add_paragraph("1. What is your favorite color?")
    doc.add_paragraph("A. Red")
    doc.add_paragraph("B. Blue")
    doc.add_paragraph("2. What is your plan?")
    blocks = list(_iter_block_items(doc))
    choices = extract_mc_choices(blocks, 0)
    assert [c["text"] for c in choices] == ["Red", "Blue"]
