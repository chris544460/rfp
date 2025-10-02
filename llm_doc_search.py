from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Iterable, Iterator, List, Sequence

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from PyPDF2 import PdfReader

try:  # pragma: no cover - optional dependency
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover - pdfplumber is optional at runtime
    pdfplumber = None

from answer_composer import CompletionsClient
from prompts import read_prompt


SEARCH_PROMPT = read_prompt(
    "llm_doc_search",
    (
        "You will be given a user question and a chunk of text from an uploaded document. "
        "If the chunk contains information that helps answer the question, "
        "respond with 'YES:' followed by only the relevant excerpt. "
        "Otherwise respond with 'NO'."
    ),
)


def _normalize_cell_text(text: str) -> str:
    text = text.strip()
    if not text:
        return ""
    text = text.replace("\n", " / ")
    return re.sub(r"\s+", " ", text)


def _rows_to_markdown(rows: Sequence[Sequence[str]]) -> str:
    if not rows:
        return ""
    max_cols = max((len(row) for row in rows), default=0)
    if max_cols == 0:
        return ""
    padded_rows = [list(row) + [""] * (max_cols - len(row)) for row in rows]
    header = padded_rows[0]
    separator = ["---"] * max_cols
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in padded_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _iter_docx_block_items(doc: Document) -> Iterator[Paragraph | Table]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _docx_table_to_rows(table: Table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        row_cells: List[str] = []
        for cell in row.cells:
            parts: List[str] = []
            for paragraph in cell.paragraphs:
                normalized = _normalize_cell_text(paragraph.text)
                if normalized:
                    parts.append(normalized)
            for inner in cell.tables:
                inner_rows = _docx_table_to_rows(inner)
                markdown = _rows_to_markdown(inner_rows)
                if markdown:
                    parts.append(markdown)
            row_cells.append(_normalize_cell_text(" \n ".join(parts)))
        rows.append(row_cells)
    return rows


def _extract_docx_text(path: str) -> str:
    doc = Document(path)
    blocks: List[str] = []
    table_index = 0
    for item in _iter_docx_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text:
                blocks.append(text)
        elif isinstance(item, Table):
            table_index += 1
            rows = _docx_table_to_rows(item)
            markdown = _rows_to_markdown([[ _normalize_cell_text(cell) for cell in row] for row in rows])
            markdown = markdown.strip()
            if markdown:
                blocks.append(f"[Table {table_index}]\n{markdown}")
    return "\n\n".join(blocks)


def _rows_from_pdf_table(table_rows: Sequence[Sequence[str]]) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table_rows:
        rows.append([_normalize_cell_text(str(cell) if cell is not None else "") for cell in row])
    return rows


def _extract_pdf_with_pdfplumber(path: str) -> str:
    if pdfplumber is None:  # pragma: no cover - handled by caller
        return ""
    pages_output: List[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                blocks: List[tuple[float, str]] = []
                table_regions: List[dict] = []
                for idx, table in enumerate(page.find_tables(), start=1):
                    data = table.extract()
                    if not data:
                        continue
                    markdown = _rows_to_markdown(_rows_from_pdf_table(data))
                    if not markdown:
                        continue
                    table_regions.append(
                        {
                            "top": table.bbox[1],
                            "bottom": table.bbox[3],
                            "left": table.bbox[0],
                            "right": table.bbox[2],
                            "text": f"[Table {idx} | Page {page_number}]\n{markdown}",
                        }
                    )
                text_lines = page.extract_text_lines() or []
                for line in text_lines:
                    text = (line.get("text") or "").strip()
                    if not text:
                        continue
                    midpoint = (line.get("top", 0.0) + line.get("bottom", 0.0)) / 2
                    baseline_x = (line.get("x0", 0.0) + line.get("x1", 0.0)) / 2
                    if any(
                        region["top"] <= midpoint <= region["bottom"]
                        and region["left"] <= baseline_x <= region["right"]
                        for region in table_regions
                    ):
                        continue
                    blocks.append((line.get("top", 0.0), text))
                blocks.extend((region["top"], region["text"]) for region in table_regions)
                blocks.sort(key=lambda item: item[0])
                if blocks:
                    page_text = "\n".join(item[1] for item in blocks)
                    pages_output.append(f"[Page {page_number}]\n{page_text}")
    except Exception:
        return ""
    return "\n\n".join(pages_output)


def _extract_pdf_with_pypdf(path: str) -> str:
    reader = PdfReader(path)
    parts: List[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt:
            parts.append(txt)
    return "\n".join(parts)


def _extract_text_from_doc(path: str) -> str:
    """Extract plain text from a .docx or .pdf file."""
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return _extract_docx_text(path)
    if ext == ".pdf":
        text = _extract_pdf_with_pdfplumber(path) if pdfplumber else ""
        if text.strip():
            return text
        return _extract_pdf_with_pypdf(path)
    raise ValueError(f"Unsupported file type: {path}")


def _iter_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> Iterable[str]:
    words = text.split()
    step = max(1, chunk_size - overlap)
    for i in range(0, len(words), step):
        yield " ".join(words[i:i + chunk_size])


def search_uploaded_docs(
    question: str,
    doc_paths: List[str],
    llm: CompletionsClient,
    chunk_size: int = 500,
    overlap: int = 50,
    context_pad: int = 50,
) -> List[Dict]:
    """Return LLM-retrieved snippets from uploaded documents.

    Each hit mirrors the structure returned by the vector search module:
    {"text": snippet, "meta": {"source": path}, "cosine": 1.0}
    """
    hits: List[Dict] = []
    for path in doc_paths:
        try:
            text = _extract_text_from_doc(path)
        except Exception:
            continue
        for chunk in _iter_chunks(text, chunk_size=chunk_size, overlap=overlap):
            prompt = (
                f"{SEARCH_PROMPT}\n\nQuestion: {question}\n\nChunk:\n{chunk}\n"
            )
            raw = llm.get_completion(prompt)
            content = raw[0] if isinstance(raw, tuple) else raw
            if not isinstance(content, str):
                continue
            reply = content.strip()
            if reply.upper().startswith("YES:"):
                snippet = reply[4:].strip()
                lower_chunk = chunk.lower()
                idx = lower_chunk.find(snippet.lower())
                if idx >= 0:
                    start = max(0, idx - context_pad)
                    end = min(len(chunk), idx + len(snippet) + context_pad)
                    snippet = chunk[start:end]
                hits.append({
                    "text": snippet,
                    "meta": {"source": str(path)},
                    "cosine": 1.0,
                    "origin": "uploaded_doc",
                })
    return hits
