#!/usr/bin/env python3
"""
cli_app.py — RFP Responder CLI

Modes by input type:
- XLSX/XLS: detect questions from a worksheet template and write answers back to a new workbook
- DOCX (default): detect question "slots" in the template and apply answers into the same layout
- PDF/TXT (and DOCX with --docx-as-text): extract questions as text, answer, and build a Q/A Word report

Dependencies in this repo:
- answer/answer_composer.CompletionsClient
- search/vector_search.search
- rfp_docx_slot_finder.extract_slots_from_docx
- rfp_docx_apply_answers.apply_answers_to_docx
"""

from __future__ import annotations

import sys
import io
import re
import os
import json
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Tuple, Dict, Any, Callable

import PyPDF2
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from openpyxl import load_workbook

# Save current sys.path
_original_sys_path = sys.path.copy()

# Expand ~ to the home directory
rfp_dir = os.path.expanduser("~/derivs-tool/rfp-ai-tool")
sys.path.insert(0, rfp_dir)

from input_file_reader.interpreter_sheet import collect_non_empty_cells
from answer.answer_composer import CompletionsClient
from search.vector_search import search

# Restore sys.path
sys.path = _original_sys_path

BASE_DIR = Path(__file__).parent

from rfp_docx_slot_finder import extract_slots_from_docx
from rfp_docx_apply_answers import apply_answers_to_docx
from rfp_xlsx_apply_answers import write_excel_answers
from rfp_xlsx_slot_finder import ask_sheet_schema
from qa_core import answer_question


def load_prompts(base: Path) -> Dict[str, str]:
    return {
        "extract_questions": (base / "prompts" / "extract_questions.txt").read_text(encoding="utf-8"),
        "answer_search_context": (base / "prompts" / "answer_search_context.txt").read_text(encoding="utf-8"),
        "answer_llm": (base / "prompts" / "answer_llm_template.txt").read_text(encoding="utf-8"),
    }


PROMPTS = load_prompts(BASE_DIR)

PRESET_INSTRUCTIONS = {
    "short": "Answer briefly in 1–2 sentences.",
    "medium": "Answer in one concise paragraph.",
    "long": "Answer in detail (up to one page).",
}


def load_input_text(path: Optional[str]) -> str:
    if path is None:
        print("[DEBUG] Reading text from stdin")
        return sys.stdin.read()
    p = Path(path)
    if not p.exists():
        print(f"[ERROR] File not found: {p}", file=sys.stderr)
        sys.exit(1)

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        out = []
        with p.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                out.append(page.extract_text() or "")
        return "\n".join(out)
    if suffix in (".docx", ".doc"):
        doc = Document(p)
        return "\n".join(par.text for par in doc.paragraphs)
    return p.read_text(encoding="utf-8")


def extract_questions(text: str, lm: CompletionsClient) -> List[str]:
    tpl = PROMPTS["extract_questions"]
    prompt = tpl.format(text=text)
    print("[DEBUG] Extracting questions via LLM")
    resp = lm.get_completion(prompt).strip()
    lines = resp.splitlines()
    out: List[str] = []
    for line in lines:
        m = re.match(r"^\s*\(?(\d+)[\).\-\:]\s*(.+?)\s*$", line)
        if m:
            out.append(m.group(2).strip())
    print(f"[DEBUG] Extracted {len(out)} questions")
    return out


def build_docx(
    questions: List[str],
    answers: List[str],
    comments: List[List[Tuple[str, str, str, float, str]]],
    include_comments: bool,
) -> bytes:
    print("[DEBUG] Building DOCX")
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    def _ensure_comments_part(document: Document):
        try:
            return document.part._comments_part
        except AttributeError:
            return document.part._add_comments_part()

    def _make_r(text: str, bold: bool = False):
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        return r

    com_part = None
    cid = 0

    def add_comment(run, lbl, src, snippet, score, date_str):
        nonlocal com_part, cid
        if com_part is None:
            com_part = _ensure_comments_part(doc)
        # anchor + details
        p = OxmlElement("w:p")
        p.append(_make_r("Citation:", True))
        p.append(_make_r(f"[{lbl}]"))
        com_part.element.append(p)
        cid += 1

    for q, ans, cmts in zip(questions, answers, comments):
        pq = doc.add_paragraph()
        rq = pq.add_run("Q: ")
        rq.bold = True
        pq.add_run(q)

        pa = doc.add_paragraph()
        ra = pa.add_run("A: ")
        ra.bold = True
        if include_comments:
            parts = re.split(r"(\[\d+\])", ans)
            for seg in parts:
                m = re.match(r"^\[(\d+)\]$", seg)
                if m:
                    idx = int(m.group(1)) - 1
                    if 0 <= idx < len(cmts):
                        lbl, src, snippet, score, date_str = cmts[idx]
                        run = pa.add_run(f"[{idx+1}]")
                        add_comment(run, lbl, src, snippet, score, date_str)
                else:
                    pa.add_run(re.sub(r"\[(\d+)\]", "", seg))
        else:
            pa.add_run(re.sub(r"\[(\d+)\]", "", ans))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_generator(
    *,
    search_mode: str,
    fund: Optional[str],
    k: int,
    length: Optional[str],
    approx_words: Optional[int],
    min_confidence: float,
    include_citations_in_text: bool,
    lm: CompletionsClient,
) -> Callable[[str], str]:
    def gen(question: str) -> str:
        ans, cmts = answer_question(
            question,
            search_mode,
            fund,
            k,
            length,
            approx_words,
            min_confidence,
            lm,
        )
        if not include_citations_in_text:
            ans = re.sub(r"\[(\d+)\]", "", ans)
        return ans

    return gen


def main():
    parser = argparse.ArgumentParser(description="RFP Responder CLI")
    parser.add_argument("input_file", help="PDF/DOCX/TXT/XLSX/XLS")
    parser.add_argument("--fund", required=False, help="Fund tag filter")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--length", choices=["short", "medium", "long"], help="Preset length")
    parser.add_argument("--approx_words", type=int, help="Approximate word count")
    parser.add_argument("--include_citations", action="store_true", help="Include [n] citations")
    parser.add_argument("--search_mode", choices=["answer", "question", "blend", "dual", "both"], default="dual")
    parser.add_argument("--lm_model", choices=["gpt-4o", "gpt-4.1", "o4-mini"], default="gpt-4o")
    parser.add_argument("--max_hits", type=int, default=6, help="Hits per question")
    parser.add_argument("--min_confidence", type=float, default=0.0, help="Min confidence threshold")
    parser.add_argument("-o", "--output", help="Output file path")
    parser.add_argument("--docx-as-text", action="store_true", help="Treat docx like free text")
    parser.add_argument("--docx-write-mode", choices=["fill", "append"], default="fill")
    parser.add_argument("--slots", help="If provided, also save detected slots JSON here")
    args = parser.parse_args()

    infile = Path(args.input_file)
    if not infile.exists():
        print(f"[ERROR] Input not found: {infile}", file=sys.stderr)
        sys.exit(1)

    suffix = infile.suffix.lower()
    lm = CompletionsClient(model=args.lm_model)

    # Excel flow
    if suffix in (".xlsx", ".xls"):
        print(f"[DEBUG] Excel input detected: {infile}")
        cells = collect_non_empty_cells(infile)
        print(f"[DEBUG] Found {len(cells)} non-empty cells")
        schema = ask_sheet_schema(infile)
        print(f"[DEBUG] Schema entries: {len(schema)}")

        answers: List[object] = []
        for i, entry in enumerate(schema, start=1):
            qtext = entry.get("question_text", "").strip()
            print(f"[DEBUG] Answering question {i}/{len(schema)} via gen")
            ans = lm.get_completion(qtext)
            answers.append(ans)

        out_path = Path(args.output or infile.with_name(infile.stem + "_answered.xlsx"))
        write_excel_answers(schema, answers, out_path)
        print(f"[DEBUG] Wrote filled Excel to {out_path}")
        sys.exit(0)

    # DOCX slot mode
    if suffix == ".docx" and not args.docx_as_text:
        print(f"[DEBUG] DOCX template detected: {infile}")
        slots_payload = extract_slots_from_docx(infile)
        if args.slots:
            Path(args.slots).write_text(json.dumps(slots_payload, indent=2), encoding="utf-8")
            slots_path = str(Path(args.slots))
        else:
            import tempfile
            fd, tmp = tempfile.mkstemp(prefix="slots_", suffix=".json")
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(slots_payload, f, indent=2)
            slots_path = tmp

        out_path = Path(args.output or infile.with_name(infile.stem + "_answered.docx"))
        summary = apply_answers_to_docx(
            docx_path=str(infile),
            slots_json_path=slots_path,
            answers_json_path=None,
            out_path=str(out_path),
            mode=args.docx_write_mode,
            generator=_make_docx_generator(
                search_mode=args.search_mode,
                fund=args.fund,
                k=args.max_hits,
                length=args.length,
                approx_words=args.approx_words,
                min_confidence=args.min_confidence,
                include_citations_in_text=args.include_citations,
                lm=lm,
            ),
            gen_name="cli_app:rag_gen",
        )
        print(f"[DEBUG] Apply summary: {summary}")
        print(f"[OK] Wrote {out_path}")
        if not args.slots:
            try:
                os.unlink(slots_path)
            except Exception:
                pass
        sys.exit(0)

    # Free-text flow
    raw = load_input_text(str(infile))
    if not raw.strip():
        print("[ERROR] No text extracted", file=sys.stderr)
        sys.exit(1)
    questions = extract_questions(raw, lm)
    if not questions:
        print("[ERROR] No questions found", file=sys.stderr)
        sys.exit(1)

    answers: List[str] = []
    comments: List[List[Tuple[str, str, str, float, str]]] = []
    for i, q in enumerate(questions, start=1):
        print(f"[DEBUG] Answering question {i}/{len(questions)}")
        ans, cmts = answer_question(
            q,
            args.search_mode,
            args.fund,
            args.max_hits,
            args.length,
            args.approx_words,
            args.min_confidence,
            lm,
        )
        answers.append(ans)
        comments.append(cmts)

    qa_doc = build_docx(questions, answers, comments, include_comments=not args.slots)
    out_path = Path(args.output or infile.with_name(infile.stem + "_answered.docx"))
    out_path.write_bytes(qa_doc)
    print(f"[DEBUG] Wrote Q/A Word report to {out_path}")


def json_dump(obj: Any) -> str:
    import json
    return json.dumps(obj, indent=2, ensure_ascii=False)


if __name__ == "__main__":
    main()
