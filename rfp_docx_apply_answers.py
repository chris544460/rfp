#!/usr/bin/env python3
# rfp_docx_apply_answers.py
# Apply answers into a DOCX according to slots.json produced by rfp_docx_slot_finder.py

import argparse, json, os, sys, re, asyncio
import importlib
from types import ModuleType
from typing import List, Union, Optional, Dict, Tuple, Callable
#
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

# ---------------------------- Debug helpers ----------------------------

DEBUG = False

def dbg(msg: str):
    if DEBUG:
        print(f"[APPLY-DEBUG] {msg}")

# ---------------------------- DOC iteration ----------------------------

def iter_block_items(doc: docx.document.Document) -> List[Union[Paragraph, Table]]:
    """
    Return a list of blocks (Paragraph or Table) in document flow order.
    Matches the approach used in the detector (paragraphs and tables interleaved).
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    blocks: List[Union[Paragraph, Table]] = []
    parent = doc.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, doc))
    return blocks

def build_indexes(doc: docx.document.Document) -> Tuple[
    List[Union[Paragraph, Table]],           # blocks
    List[Paragraph],                          # paragraphs_only
    Dict[int, int],                           # block_index -> paragraph_index (if block is paragraph)
    Dict[int, int]                            # block_index -> table_index (if block is table)
]:
    """
    Build convenient indexes over the doc content.
    """
    blocks = iter_block_items(doc)
    paragraphs: List[Paragraph] = []
    block_to_para: Dict[int, int] = {}
    block_to_table: Dict[int, int] = {}
    running_table_index = 0
    for bi, b in enumerate(blocks):
        if isinstance(b, Paragraph):
            block_to_para[bi] = len(paragraphs)
            paragraphs.append(b)
        elif isinstance(b, Table):
            block_to_table[bi] = running_table_index
            running_table_index += 1
    return blocks, paragraphs, block_to_para, block_to_table

# ---------------------------- Utilities ----------------------------

_BLANK_RE = re.compile(r"_+\s*$")

_CHECKBOX_CHARS = "☐☑☒□■✓✔✗✘"

# Matches bracketed citation numbers like "[1]" used in answers
_CITATION_RE = re.compile(r"\[(\d+)\]")

def is_blank_para(p: Paragraph) -> bool:
    """
    Heuristic 'blank' detection similar to detector:
    empty, underscores, or brackets placeholders.
    """
    t = (p.text or "").strip()
    if t == "":
        return True
    if _BLANK_RE.match(t):
        return True
    if re.fullmatch(r"\[(?:insert|enter|provide)[^\]]*\]", t.lower()):
        return True
    # if any run is underlined with no visible text
    try:
        if any(r.text and r.underline for r in p.runs) and len(t.replace("_","").strip()) == 0:
            return True
    except Exception:
        pass
    return False

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """
    Insert a new paragraph XML immediately after the given paragraph.
    Returns a python-docx Paragraph wrapper for the new element.
    """
    new_p_elm = OxmlElement("w:p")
    paragraph._element.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)
    if text:
        new_p.add_run(text)
    return new_p

def normalize_question(q: str) -> str:
    return " ".join((q or "").strip().lower().split())


def _normalize_citations(raw: object) -> Dict[str, str]:
    """Return a {citation_number: snippet_text} mapping from ``raw``.

    ``raw`` may already be a mapping or a list of citation objects. Each citation
    object is expected to contain the snippet text either as the value itself or
    under a common key like ``text`` or ``snippet``. Unknown structures yield an
    empty mapping.
    """
    if not raw:
        return {}
    result: Dict[str, str] = {}
    if isinstance(raw, dict):
        items = raw.items()
    elif isinstance(raw, list):  # allow list of citation objects
        items = []
        for i, item in enumerate(raw, 1):
            key = getattr(item, "get", lambda *_: None)("id") or getattr(item, "get", lambda *_: None)("num") or i
            items.append((key, item))
    else:
        return {}
    for key, val in items:
        if isinstance(val, dict):
            snippet = val.get("text") or val.get("snippet") or val.get("source_text") or val.get("content") or ""
        else:
            snippet = str(val)
        result[str(key)] = str(snippet)
    return result


def _add_text_with_citations(paragraph: Paragraph, text: str, citations: Dict[object, str]) -> None:
    """Write text into ``paragraph`` adding Word comments for citation markers.

    ``citations`` is a mapping from citation number to the snippet text that
    should appear in the comment. Each occurrence of "[n]" in ``text`` results
    in a comment attached to the run containing that marker when ``n`` is found
    in ``citations``.
    """
    doc = paragraph.part.document
    parts = text.split("\n")
    for li, line in enumerate(parts):
        pos = 0
        for match in _CITATION_RE.finditer(line):
            if match.start() > pos:
                paragraph.add_run(line[pos:match.start()])
            num = int(match.group(1))
            run = paragraph.add_run(match.group(0))
            snippet = citations.get(num) or citations.get(str(num))
            if snippet:
                doc.add_comment(run, str(snippet))
            pos = match.end()
        if pos < len(line):
            paragraph.add_run(line[pos:])
        if li < len(parts) - 1:
            paragraph.add_run().add_break()

# ---------------------------- Answers loader ----------------------------

def load_answers(answers_path: str) -> Tuple[Dict[str, object], Dict[str, object]]:
    """
    Load answers from JSON.

    Accepted formats:
    1) Dict with explicit sections:
       {
         "by_id": {"slot_abc": "answer" | {"text": "...", "citations": {...}}, ...},
         "by_question": {"what’s 1+1?": "2", ...}
       }

    2) List of objects:
       [{"slot_id":"slot_abc","answer":"..."},
        {"question_text":"...","answer":{"text":"...","citations":{...}}}]

    3) Flat dict:
       Keys matching slot IDs go to by_id; others are treated as question_text.
       Values may be plain strings or objects with ``text``/``citations``.
    """
    with open(answers_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # answers may be plain strings or dicts containing citation metadata
    by_id: Dict[str, object] = {}
    by_q: Dict[str, object] = {}

    if isinstance(data, dict):
        if "by_id" in data or "by_question" in data:
            for k, v in (data.get("by_id") or {}).items():
                by_id[str(k)] = v
            for k, v in (data.get("by_question") or {}).items():
                by_q[normalize_question(k)] = v
        else:
            # flat dict
            for k, v in data.items():
                kstr = str(k)
                if kstr.startswith("slot_"):
                    by_id[kstr] = v
                else:
                    by_q[normalize_question(kstr)] = v
    elif isinstance(data, list):
        for item in data:
            if not isinstance(item, dict):
                continue
            if "slot_id" in item:
                by_id[str(item["slot_id"])] = item.get("answer", "")
            elif "question_text" in item:
                by_q[normalize_question(str(item["question_text"]))] = item.get("answer", "")
    else:
        raise ValueError("Unsupported answers JSON structure.")

    return by_id, by_q

# ---------------------------- Locator resolution ----------------------------

def resolve_anchor_paragraph(
    doc: docx.document.Document,
    blocks: List[Union[Paragraph, Table]],
    paragraphs: List[Paragraph],
    block_to_para: Dict[int, int],
    locator: Dict[str, object],
    meta: Optional[Dict[str, object]]
) -> Optional[Paragraph]:
    """
    Resolve the anchor paragraph for a locator that references a question area.

    Supports two common cases for locator['type'] in slots.json:
      - 'paragraph_after'  (anchor = question paragraph)
      - 'paragraph'        (anchor = the answer paragraph itself; caller can use it directly)
    """
    ltype = str(locator.get("type", ""))
    p_idx = locator.get("paragraph_index")
    p_idx = int(p_idx) if p_idx is not None else None

    # For 'paragraph' locators, 'paragraph_index' is already a paragraph-only index (legacy heuristics path).
    if ltype == "paragraph":
        if p_idx is None:
            return None
        if 0 <= p_idx < len(paragraphs):
            return paragraphs[p_idx]
        return None

    # For 'paragraph_after', interpret paragraph_index as:
    #   (a) Global block index of the question (two_stage / llm_rich),
    #   (b) Or paragraph-only index (legacy).
    if ltype == "paragraph_after":
        # Prefer explicit q_block from meta when present
        qb = None
        if isinstance(meta, dict) and "q_block" in meta:
            try:
                qb = int(meta["q_block"])
            except Exception:
                qb = None

        # Option 1: meta.q_block is a valid global block index
        if qb is not None and 0 <= qb < len(blocks) and isinstance(blocks[qb], Paragraph):
            return blocks[qb]  # type: ignore

        # Option 2: paragraph_index is a global block index pointing at a Paragraph
        if p_idx is not None and 0 <= p_idx < len(blocks) and isinstance(blocks[p_idx], Paragraph):
            return blocks[p_idx]  # type: ignore

        # Option 3: paragraph_index is already a paragraph-only index
        if p_idx is not None and 0 <= p_idx < len(paragraphs):
            return paragraphs[p_idx]

    return None

def get_target_paragraph_after_anchor(
    blocks: List[Union[Paragraph, Table]],
    block_to_para: Dict[int, int],
    paragraphs: List[Paragraph],
    anchor_para: Paragraph,
    offset: int
) -> Paragraph:
    """
    Find the paragraph that is the N-th paragraph AFTER the anchor, scanning through blocks.
    If there aren't enough following paragraphs, insert new paragraphs directly after the anchor
    until we reach the desired offset.
    """
    # Locate anchor in blocks & paragraph index
    anchor_para_index = None
    anchor_block_index = None
    for bi, b in enumerate(blocks):
        if isinstance(b, Paragraph) and b is anchor_para:
            anchor_block_index = bi
            anchor_para_index = block_to_para[bi]
            break
    if anchor_block_index is None or anchor_para_index is None:
        # Fallback: use linear paragraph list
        anchor_para_index = paragraphs.index(anchor_para)

    # Collect existing subsequent paragraphs in flow order
    subsequent_paras: List[Paragraph] = []
    if anchor_block_index is not None:
        for b in blocks[anchor_block_index + 1:]:
            if isinstance(b, Paragraph):
                subsequent_paras.append(b)
            # Note: tables are ignored when counting "paragraph_after" offset
    else:
        # No block index found; use paragraph list instead
        subsequent_paras = paragraphs[anchor_para_index + 1:]

    # If we already have enough paragraphs ahead, return the offset-th
    if len(subsequent_paras) >= offset:
        return subsequent_paras[offset - 1]

    # Else, insert new paragraphs after the anchor (or last found), enough to meet the offset
    needed = offset - len(subsequent_paras)
    dbg(f"Not enough following paragraphs: need to insert {needed} after anchor")
    last = anchor_para if not subsequent_paras else subsequent_paras[-1]
    created: Paragraph = last
    for _ in range(needed):
        created = insert_paragraph_after(created, "")
    return created

# ---------------------------- Apply operations ----------------------------

def apply_to_paragraph(target: Paragraph, answer: object, mode: str = "fill") -> None:
    """Write answer into a paragraph, supporting citation comments.

    ``answer`` may be a plain string or a dict with ``text`` and ``citations`` fields.
    """
    if isinstance(answer, dict):
        answer_text = str(answer.get("text", ""))
        citations = _normalize_citations(answer.get("citations"))
    else:
        answer_text = str(answer)
        citations = {}

    existing = target.text or ""
    if mode == "replace":
        target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        return
    if mode == "append":
        if existing:
            target.add_run("\n")
        else:
            target.text = ""
        _add_text_with_citations(target, answer_text, citations)
        return
    # fill (default)
    if is_blank_para(target) or not existing.strip():
        target.text = ""
        _add_text_with_citations(target, answer_text, citations)
    else:
        new_p = insert_paragraph_after(target, "")
        _add_text_with_citations(new_p, answer_text, citations)
        dbg("Target paragraph not blank; appended answer in a new paragraph below.")

def apply_to_table_cell(tbl: Table, row: int, col: int, answer: object, mode: str = "fill") -> None:
    """Write into a table cell, supporting citation comments."""
    try:
        cell = tbl.cell(row, col)
    except Exception:
        return

    if isinstance(answer, dict):
        answer_text = str(answer.get("text", ""))
        citations = _normalize_citations(answer.get("citations"))
    else:
        answer_text = str(answer)
        citations = {}

    current = cell.text or ""
    if mode == "replace":
        cell.text = ""
        p = cell.paragraphs[0]
        _add_text_with_citations(p, answer_text, citations)
        return

    if current.strip():
        p = cell.paragraphs[-1]
        p.add_run().add_break()
        _add_text_with_citations(p, answer_text, citations)
    else:
        cell.text = ""
        p = cell.paragraphs[0]
        _add_text_with_citations(p, answer_text, citations)

def mark_multiple_choice(blocks: List[Union[Paragraph, Table]], choices_meta: List[Dict[str, object]], index: int, style: Optional[str] = None) -> None:
    """Mark a selected multiple-choice option in-place."""
    if not isinstance(choices_meta, list):
        return
    if index is None or not (0 <= index < len(choices_meta)):
        return
    meta = choices_meta[index]
    b_idx = int(meta.get("block_index", -1))
    if not (0 <= b_idx < len(blocks)):
        return
    para = blocks[b_idx]
    if not isinstance(para, Paragraph):
        return
    prefix = str(meta.get("prefix", ""))
    text = para.text or ""
    # auto style detection
    if style in (None, "", "auto"):
        if any(ch in prefix for ch in _CHECKBOX_CHARS):
            style = "checkbox"
        elif prefix.strip() in ("()", "[]"):
            style = "fill"
        else:
            style = "highlight"
    if style == "checkbox" and any(ch in prefix for ch in _CHECKBOX_CHARS):
        para.text = re.sub(rf"[{_CHECKBOX_CHARS}]", "☑", text, count=1)
    elif style == "fill" and prefix.strip() in ("()", "[]"):
        mark = prefix[0] + "X" + prefix[1]
        para.text = text.replace(prefix, mark, 1)
    elif style == "highlight":
        for run in para.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        para.text = "X " + text

# ---------------------------- Main application flow ----------------------------

def apply_answers_to_docx(
    docx_path: str,
    slots_json_path: str,
    answers_json_path: str,
    out_path: str,
    mode: str = "fill",
    generator: Optional[Callable[..., object]] = None,
    gen_name: str = ""
) -> Dict[str, int]:
    """
    Apply answers into the DOCX and save to out_path.
    Returns a small summary dict with counts.
    """
    # Load inputs
    with open(slots_json_path, "r", encoding="utf-8") as f:
        slots_payload = json.load(f)

    by_id, by_q = ({}, {})
    if answers_json_path and answers_json_path != "-" and os.path.isfile(answers_json_path):
        by_id, by_q = load_answers(answers_json_path)
        dbg(f"Answers loaded: by_id={len(by_id)}, by_question={len(by_q)}")
    else:
        dbg("No answers file provided; relying solely on generator (if any)")

    doc = docx.Document(docx_path)
    blocks, paragraphs, block_to_para, block_to_table = build_indexes(doc)

    applied = 0
    skipped_no_answer = 0
    skipped_bad_locator = 0
    skipped_table_oob = 0
    generated = 0

    slots = (slots_payload or {}).get("slots", [])

    # First pass: resolve answers from file or schedule generation
    answers: Dict[str, Optional[object]] = {}
    to_generate: List[Tuple[str, str, dict]] = []
    for s in slots:
        sid = s.get("id", "")
        question_text = (s.get("question_text") or "").strip()
        meta = s.get("meta") or {}
        answer = None
        if sid in by_id:
            answer = by_id[sid]
        else:
            key = normalize_question(question_text)
            if key in by_q:
                answer = by_q[key]
        if answer is None and generator is not None:
            kwargs = {}
            if s.get("answer_type") == "multiple_choice":
                choice_meta = meta.get("choices", [])
                kwargs["choices"] = [c.get("text") if isinstance(c, dict) else str(c) for c in choice_meta]
                kwargs["choice_meta"] = choice_meta
            to_generate.append((sid, question_text, kwargs))
        answers[sid] = answer

    # Run generator concurrently for all missing answers
    if to_generate:
        async def run_all() -> List[Tuple[str, Optional[object]]]:
            async def worker(sid: str, q: str, kwargs: dict):
                try:
                    ans = await asyncio.to_thread(generator, q, **kwargs)
                    dbg(f"Generated answer via {gen_name} for slot {sid}: {ans}")
                    return sid, ans
                except Exception as e:
                    dbg(f"Generator error for question '{q}': {e}")
                    return sid, None
            tasks = [asyncio.create_task(worker(sid, q, kw)) for sid, q, kw in to_generate]
            return await asyncio.gather(*tasks)

        for sid, ans in asyncio.run(run_all()):
            if ans is not None:
                answers[sid] = ans
                generated += 1

    # Second pass: apply answers
    for s in slots:
        sid = s.get("id", "")
        question_text = (s.get("question_text") or "").strip()
        locator = s.get("answer_locator") or {}
        ltype = str(locator.get("type", ""))
        meta = s.get("meta") or {}

        answer = answers.get(sid)
        if answer is None:
            dbg(f"NO ANSWER for slot {sid!r} / question '{question_text}' — skipping")
            skipped_no_answer += 1
            continue

        # Special handling for multiple-choice selections
        choice_meta = meta.get("choices", [])
        if s.get("answer_type") == "multiple_choice" and choice_meta:
            choice_texts = [c.get("text") if isinstance(c, dict) else str(c) for c in choice_meta]
            idx = None
            style = None
            if isinstance(answer, dict):
                idx = answer.get("choice_index")
                style = answer.get("style")
            else:
                try:
                    idx = choice_texts.index(str(answer).strip())
                except ValueError:
                    idx = None
            if idx is not None:
                try:
                    mark_multiple_choice(blocks, choice_meta, int(idx), style)
                    applied += 1
                    dbg("  -> marked choice in-place")
                except Exception as e:
                    dbg(f"  -> error marking multiple choice: {e}")
                    skipped_bad_locator += 1
            else:
                dbg("  -> could not resolve selected choice index")
                skipped_bad_locator += 1
            continue

        dbg(f"Applying answer for slot {sid} (type={ltype})")

        try:
            if ltype == "table_cell":
                t_index = locator.get("table_index")
                row = locator.get("row", 0)
                col = locator.get("col", 0)
                if t_index is None:
                    dbg("  -> bad locator: missing table_index")
                    skipped_bad_locator += 1
                    continue
                t_index = int(t_index)
                row = int(row)
                col = int(col)
                # Find the table by overall order
                if 0 <= t_index < len(doc.tables):
                    tbl = doc.tables[t_index]
                    apply_to_table_cell(tbl, row, col, answer, mode=mode)
                    applied += 1
                    dbg(f"  -> wrote into table[{t_index}] cell({row},{col})")
                else:
                    dbg(f"  -> table_index {t_index} out of bounds; have {len(doc.tables)} tables")
                    skipped_table_oob += 1

            elif ltype in ("paragraph_after", "paragraph"):
                # Resolve anchor (question paragraph) or direct target paragraph
                anchor = resolve_anchor_paragraph(doc, blocks, paragraphs, block_to_para, locator, meta)
                if anchor is None:
                    dbg("  -> could not resolve anchor/target paragraph")
                    skipped_bad_locator += 1
                    continue

                if ltype == "paragraph_after":
                    offset = int(locator.get("offset", 1) or 1)
                    target = get_target_paragraph_after_anchor(blocks, block_to_para, paragraphs, anchor, offset)
                else:
                    # 'paragraph' locators already resolved to the paragraph itself
                    target = anchor

                apply_to_paragraph(target, answer, mode=mode)
                applied += 1
                dbg("  -> wrote into paragraph")

            else:
                dbg(f"  -> unsupported locator type: {ltype}")
                skipped_bad_locator += 1

        except Exception as e:
            dbg(f"  -> error while applying slot {sid}: {e}")
            # treat as bad locator for summary
            skipped_bad_locator += 1

    # Save document
    doc.save(out_path)
    print(f"Wrote {out_path}")

    return {
        "applied": applied,
        "skipped_no_answer": skipped_no_answer,
        "skipped_bad_locator": skipped_bad_locator,
        "skipped_table_oob": skipped_table_oob,
        "total_slots": len(slots),
        "generated": generated,
    }

# ---------------------------- CLI ----------------------------

def main():
    ap = argparse.ArgumentParser(description="Apply answers into a DOCX using slots.json")
    ap.add_argument("docx_path", help="Path to the original .docx")
    ap.add_argument("slots_json", help="Path to slots.json produced by the detector")
    ap.add_argument("answers_json", nargs="?", default="", help="Path to answers.json (optional if using --generate)")
    ap.add_argument("-o", "--out", required=True, help="Path to write updated .docx")
    ap.add_argument("--mode", choices=["replace", "append", "fill"], default="fill",
                    help="Write mode for paragraphs/cells (default: fill)")
    ap.add_argument("--debug", action="store_true", help="Verbose debug logging")
    ap.add_argument("--generate", metavar="MODULE:FUNC", help="Dynamically generate answers by calling given function for each question (e.g. ai_gen:make_answer)")
    if len(sys.argv) == 1:
        ap.print_help()
        sys.exit(1)
    args = ap.parse_args()

    global DEBUG
    DEBUG = args.debug
    if DEBUG:
        print("### APPLY DEBUG MODE ON ###")

    # Validate inputs
    required_paths = [args.docx_path, args.slots_json]
    for p in required_paths:
        if not os.path.isfile(p):
            print(f"Error: '{p}' does not exist.", file=sys.stderr)
            sys.exit(1)
    if args.answers_json and args.answers_json != "-" and not os.path.isfile(args.answers_json):
        if not args.generate:
            print(f"Error: answers file '{args.answers_json}' does not exist and no --generate specified.", file=sys.stderr)
            sys.exit(1)
        else:
            if DEBUG:
                print(f"Warning: answers file '{args.answers_json}' not found; proceeding with generator only.")

    gen_callable = None
    gen_name = ""
    if args.generate:
        if ":" not in args.generate:
            print("Error: --generate requires MODULE:FUNC", file=sys.stderr)
            sys.exit(1)
        mod_name, func_name = args.generate.split(":", 1)
        try:
            module: ModuleType = importlib.import_module(mod_name)
            gen_callable = getattr(module, func_name)
            if not callable(gen_callable):
                raise AttributeError
            gen_name = args.generate
            if DEBUG:
                print(f"Loaded generator function {gen_name}")
        except Exception as e:
            print(f"Error: failed to load generator function {args.generate}: {e}", file=sys.stderr)
            sys.exit(1)

    try:
        summary = apply_answers_to_docx(
            args.docx_path,
            args.slots_json,
            args.answers_json,
            args.out,
            mode=args.mode,
            generator=gen_callable,
            gen_name=gen_name
        )
    except Exception as e:
        print(f"Error: failed to apply answers: {e}", file=sys.stderr)
        sys.exit(1)

    if DEBUG:
        print("--- APPLY SUMMARY ---")
        for k, v in summary.items():
            print(f"{k}: {v}")

if __name__ == "__main__":
    main()
